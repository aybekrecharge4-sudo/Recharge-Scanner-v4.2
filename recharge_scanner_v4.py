"""
RECHARGE.COM OPPORTUNITY SCANNER v5.0
16 data sources | Google Search grounding | Composite scoring | 4-pass AI
Single-file script -> HTML dashboard + Word + Email + GitHub Pages
Works in: Google Colab, GitHub Actions, local Python

# =============================================================================
# TABLE OF CONTENTS
# =============================================================================
# SECTION 0  - IMPORTS & CONFIG
# SECTION 1  - DEPENDENCIES
# SECTION 2  - DATA MODELS
# SECTION 3  - UTILITIES (HTTP, fuzzy match, scoring, history)
# SECTION 4  - DATA FETCHERS (16 sources)
# SECTION 5  - CONCURRENT ORCHESTRATOR
# SECTION 6  - DEDUP & MERGE
# SECTION 7  - SCORE NORMALIZATION & COMPOSITE SCORING
# SECTION 8  - AI ANALYSIS (4-pass + newsletter)
# SECTION 9  - GOOGLE SHEETS
# SECTION 10 - HTML DASHBOARD
# SECTION 11 - WORD DOCUMENT
# SECTION 12 - EMAIL NEWSLETTER
# SECTION 13 - EVENTS CALENDAR
# SECTION 14 - MAIN ORCHESTRATOR
# =============================================================================
"""

# =============================================================================
# SECTION 0 - IMPORTS & CONFIG
# =============================================================================

import os, sys, json, re, time, logging, subprocess, html as _html
from datetime import datetime, timedelta
from collections import defaultdict, Counter
from dataclasses import dataclass, field
from typing import List, Dict, Optional
from urllib.parse import quote, urlparse
from difflib import SequenceMatcher
from concurrent.futures import ThreadPoolExecutor, as_completed
import xml.etree.ElementTree as ET

_log_fmt = "%(asctime)s [%(levelname)s] %(message)s"
logging.basicConfig(level=logging.INFO, format=_log_fmt)
log = logging.getLogger("v5")

# File handler: detailed log persisted to disk
_log_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), f"scanner_run_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}.log")
_fh = logging.FileHandler(_log_file, encoding="utf-8")
_fh.setLevel(logging.DEBUG)
_fh.setFormatter(logging.Formatter("%(asctime)s [%(levelname)-7s] %(message)s"))
log.addHandler(_fh)
log.setLevel(logging.DEBUG)
log.info(f"=== Recharge Scanner v5.0 — Log file: {_log_file} ===")

GEMINI_KEY = os.environ.get("GEMINI_API_KEY", "AIzaSyD22ZtKktZ6rRLz9AqGc4310jd9YC1n7S8")
DASHBOARD_URL = os.environ.get("DASHBOARD_URL", "")
DASH_PASSWORD = os.environ.get("DASH_PASSWORD", "")  # if set, dashboard requires password
OXYLABS_USER = os.environ.get("OXYLABS_USER", "aybek_BCWTG")
OXYLABS_PASS = os.environ.get("OXYLABS_PASS", "45119971905Aybe_")

NOW  = datetime.now()
YEAR = NOW.year
MON  = NOW.strftime("%B")
DATE = NOW.strftime("%Y-%m-%d")
TIME = NOW.strftime("%H:%M")

IS_COLAB = False
try:
    import google.colab
    IS_COLAB = True
except ImportError: pass

W = {
    # Tier 1: Real-time, high-signal sources (50%)
    "oxylabs_news": .16, "trends": .12, "news": .12, "reddit": .10,
    # Tier 2: Platform-specific actionable data (30%)
    "steam": .10, "epic": .05, "youtube": .05, "competitor": .05, "steamspy": .05,
    # Tier 3: Supplementary signals (20%)
    "cheapshark": .04, "gamerpower": .04, "gog": .03, "humble": .03,
    "anime": .03, "sitemap": .02, "wiki": .01, "freetogame": .01,
}

CONF = {1: 0.55, 2: 0.75, 3: 0.90, 4: 1.0}
CONF_DEFAULT = 1.0
FUZZ_T = 0.82
UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
      "(KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36")

# Business category mapping: every KW key -> GMG/ENT/PPM/MTU
BIZ_CATS = {
    # GMG - Gaming
    "EA Sports FC":"GMG","PlayStation":"GMG","Xbox":"GMG","Nintendo":"GMG",
    "Steam":"GMG","Fortnite":"GMG","Call of Duty":"GMG","GTA":"GMG",
    "Minecraft":"GMG","Roblox":"GMG","Valorant":"GMG","League of Legends":"GMG",
    "Genshin Impact":"GMG","Honkai":"GMG","PUBG Mobile":"GMG","Free Fire":"GMG",
    "Mobile Legends":"GMG","Discord":"ENT","Twitch":"ENT","Meta Quest":"GMG",
    "Overwatch":"GMG","World of Warcraft":"GMG","Counter-Strike":"GMG",
    "Apex Legends":"GMG","Dead by Daylight":"GMG","Esports":"GMG",
    "Destiny":"GMG","Elden Ring":"GMG","Monster Hunter":"GMG",
    # ENT - Entertainment
    "Spotify":"ENT","Netflix":"ENT","Disney Plus":"ENT",
    "Amazon Prime":"ENT","YouTube Premium":"ENT","Apple":"ENT","Crunchyroll":"ENT",
    # PPM - Prepaid Money
    "Gift Cards":"PPM","Paysafecard":"PPM","Razer Gold":"PPM","Google Play":"PPM",
    # MTU - Mobile Top Up
    "Mobile Top Up":"MTU","Lycamobile":"MTU","Lebara":"MTU",
    "T-Mobile":"MTU","Vodafone":"MTU","Orange":"MTU",
    "Gaming":"GMG",
}
BIZ_CAT_NAMES = {"GMG":"Gaming","ENT":"Entertainment","PPM":"Prepaid Money","MTU":"Mobile Top Up"}
BIZ_CAT_EMOJI = {"GMG":"\U0001F3AE","ENT":"\U0001F3AC","PPM":"\U0001F4B3","MTU":"\U0001F4F1"}
BIZ_CAT_COLORS = {"GMG":"#7c3aed","ENT":"#ec4899","PPM":"#059669","MTU":"#2563eb"}

KW = {
    "EA Sports FC":  ["EA FC","FC 26","FC 25","FIFA","FUT","TOTY","TOTS",
                      "Ultimate Team","FIFA Points","FC Points","Madden"],
    "PlayStation":   ["PlayStation","PS5","PS4","PS Plus","PlayStation Plus",
                      "PSN","PlayStation Store","PS VR","PlayStation Portal","DualSense"],
    "Xbox":          ["Xbox","Game Pass","Xbox Live","Xbox Series",
                      "Game Pass Ultimate","Game Pass Core","Xbox Cloud"],
    "Nintendo":      ["Nintendo","Switch","eShop","Nintendo Direct",
                      "Nintendo Online","Pokemon","Zelda","Mario","Switch 2","Animal Crossing"],
    "Steam":         ["Steam","Steam Deck","Steam Sale","Steam Wallet",
                      "Valve","Steam Fest","Steam Next Fest","Steam Summer","Steam Winter","Steam Spring"],
    "Fortnite":      ["Fortnite","V-Bucks","Battle Pass","Epic Games"],
    "Call of Duty":  ["Call of Duty","COD","Warzone","Modern Warfare","Black Ops","COD Points"],
    "GTA":           ["GTA","Grand Theft Auto","GTA 6","GTA Online","Shark Card","Rockstar"],
    "Minecraft":     ["Minecraft","Minecoins","Minecraft Realms"],
    "Roblox":        ["Roblox","Robux"],
    "Valorant":      ["Valorant","Valorant Points"],
    "League of Legends": ["League of Legends","LoL","Riot Points","Riot Games","Arcane"],
    "Genshin Impact":["Genshin","Primogems","Genesis Crystals","HoYoverse","Traveler","Teyvat"],
    "Honkai":        ["Honkai Star Rail","Honkai Impact","Penacony"],
    "PUBG Mobile":   ["PUBG Mobile","PUBG UC","Royale Pass","PUBG"],
    "Free Fire":     ["Free Fire","Free Fire Diamonds"],
    "Mobile Legends":["Mobile Legends","MLBB"],
    "Spotify":       ["Spotify","Spotify Premium","Spotify Wrapped"],
    "Netflix":       ["Netflix","Netflix Games","Squid Game"],
    "Disney Plus":   ["Disney Plus","Disney+","Hotstar","Star Wars","Marvel","Mandalorian"],
    "Amazon Prime":  ["Prime Video","Prime Gaming","Amazon Prime","Twitch Prime"],
    "YouTube Premium":["YouTube Premium","YouTube Music"],
    "Apple":         ["Apple TV","Apple Music","iTunes","App Store","Apple gift card","Apple Arcade"],
    "Google Play":   ["Google Play","Play Store","Google gift card"],
    "Discord":       ["Discord","Nitro","Discord Nitro"],
    "Twitch":        ["Twitch","Twitch bits","Twitch sub"],
    "Gift Cards":    ["gift card","prepaid","voucher","e-gift","top-up","recharge","digital code"],
    "Paysafecard":   ["paysafecard","Neosurf"],
    "Razer Gold":    ["Razer Gold","Karma Koin"],
    "Crunchyroll":   ["Crunchyroll","anime streaming","Funimation","One Piece","Dragon Ball",
                      "Demon Slayer","Jujutsu Kaisen","My Hero Academia","Naruto","Bleach","anime"],
    "Meta Quest":    ["Meta Quest","Quest 3","Oculus","VR","Quest Pro"],
    "Overwatch":     ["Overwatch","Overwatch 2"],
    "World of Warcraft": ["World of Warcraft","WoW","WoW Token","Blizzard","Diablo","Hearthstone"],
    "Counter-Strike":["CS2","Counter-Strike","CS:GO"],
    "Apex Legends":  ["Apex Legends","Apex Coins"],
    "Dead by Daylight": ["Dead by Daylight","DBD","Auric Cells"],
    "Esports":       ["esports","tournament","championship","Worlds","Major"],
    "Destiny":       ["Destiny 2","Bungie","Silver"],
    "Elden Ring":    ["Elden Ring","FromSoftware","Dark Souls"],
    "Monster Hunter":["Monster Hunter","Capcom","MH Wilds"],
    "Mobile Top Up":  ["mobile top up","phone credit","airtime","mobile recharge",
                       "international calling","SIM card","phone recharge",
                       "carrier top up","prepaid mobile","data bundle"],
    "Lycamobile":     ["Lycamobile","Lyca"],
    "Lebara":         ["Lebara"],
    "T-Mobile":       ["T-Mobile","T Mobile"],
    "Vodafone":       ["Vodafone"],
    "Orange":         ["Orange mobile","Orange top up"],
    "Gaming":         ["gaming","gamer","video game","game release","game update","new game"],
}

RSS_FEEDS = {
    "IGN":"https://feeds.feedburner.com/ign/all",
    "GameSpot":"https://www.gamespot.com/feeds/news/",
    "Kotaku":"https://kotaku.com/rss",
    "PC Gamer":"https://www.pcgamer.com/rss/",
    "Eurogamer":"https://www.eurogamer.net/feed",
    "The Verge":"https://www.theverge.com/rss/index.xml",
    "Polygon":"https://www.polygon.com/rss/index.xml",
    "GamesRadar":"https://www.gamesradar.com/rss/",
    "Dexerto":"https://www.dexerto.com/feed/",
    "Destructoid":"https://www.destructoid.com/feed/",
    "VG247":"https://www.vg247.com/feed/all",
    "RPS":"https://www.rockpapershotgun.com/feed",
    "Push Square":"https://www.pushsquare.com/feeds/latest",
    "Nintendo Life":"https://www.nintendolife.com/feeds/latest",
    "Pure Xbox":"https://www.purexbox.com/feeds/latest",
    "Siliconera":"https://www.siliconera.com/feed/",
    "TouchArcade":"https://toucharcade.com/feed/",
    "XboxWire":"https://news.xbox.com/en-us/feed/",
    "PlayStation Blog":"https://blog.playstation.com/feed/",
    # Gaming (8 new)
    "DualShockers":"https://www.dualshockers.com/feed/video-game-news/",
    "GameRant":"https://gamerant.com/feed/",
    "GamesIndustry.biz":"https://www.gamesindustry.biz/feed",
    "Game Developer":"https://www.gamedeveloper.com/rss.xml",
    "Game Informer":"https://gameinformer.com/news.xml",
    "Screen Rant":"https://screenrant.com/feed/",
    "Tom's Hardware":"https://www.tomshardware.com/feeds/all",
    "NME Gaming":"https://www.nme.com/gaming/feed",
    # Streaming (3 new)
    "What's On Netflix":"https://www.whats-on-netflix.com/feed/",
    "Decider":"https://decider.com/feed/",
    "The Streamable":"https://thestreamable.com/feed",
    # Payments (3 new)
    "PYMNTS":"https://www.pymnts.com/feed/",
    "Finextra":"https://www.finextra.com/rss/headlines.aspx",
    "Payment Journal":"https://www.paymentjournal.com/feed/",
    # Mobile (2 new)
    "GSMArena":"https://www.gsmarena.com/rss-news-reviews.php3",
    "PhoneArena":"https://www.phonearena.com/feed",
    # Deals (2 new)
    "Slickdeals":"https://slickdeals.net/newsearch.php?mode=frontpage&searcharea=deals&searchin=first&rss=1",
    "DealNews":"https://www.dealnews.com/rss/",
    # Gaming (round 2)
    "VGC":"https://www.videogameschronicle.com/category/news/feed/",
    "PCGamesN":"https://www.pcgamesn.com/mainrss.xml",
    "GameGeekNews":"https://gamegeeksnews.com/feed/",
    "EveryoneGaming":"https://everyonegaming.com/news/feed/",
    # Streaming & Entertainment (round 2)
    "Collider":"https://collider.com/feed/",
    "Deadline TV":"https://deadline.com/v/tv/feed/",
    "CinemaBlend":"https://www.cinemablend.com/rss",
    "ComingSoon":"https://www.comingsoon.net/feed",
    "Digital Spy":"https://www.digitalspy.com/rss/tv.xml",
    "Ready Steady Cut":"https://readysteadycut.com/feed/",
    "Moviefone":"https://www.moviefone.com/feed/",
}

YT_CHANNELS = {
    "IGN":"UCKy1dAqELo0zrOtPkf0eTMw","GameSpot":"UCbu2SsF1frCRhGHstdXZR5g",
    "PlayStation":"UC-2Y8dQb0S6DtpxNgAKoJKA","Xbox":"UCXGgrKt94gR6lmN4aN3mYTg",
    "Nintendo":"UCGIY_O-8vW4rfx98KlMkvRg","TheGameAwards":"UCMjezPLBl-5fVS0HERR6QRA",
    "SkillUp":"UCZ7AeeVbyslLM_8-nQy_8CQ","ACG":"UCK9_x1DImhU-eolIay5rb2Q",
    "DigitalFoundry":"UC9PBzalIcEQCsiIkq36PyUA","Laymen Gaming":"UCYkgPmEwIcTn_WmocdZXEcg",
    "Fextralife":"UC1ONOluGA4Hht2CsMkJOyVg",
}

SUBREDDITS = [
    "gaming","Games","pcgaming","PS5","XboxSeriesX","NintendoSwitch",
    "Steam","FortNiteBR","EASportsFC","GenshinImpact","leagueoflegends",
    "VALORANT","Roblox","spotify","netflix","GameDeals",
    "FreeGameFindings","anime","CrunchyrollPremium","NintendoSwitch2",
]

COMPETITORS = {"G2A":"https://www.g2a.com/","Eneba":"https://www.eneba.com/",
               "CDKeys":"https://www.cdkeys.com/","Kinguin":"https://www.kinguin.net/"}

SITEMAP_COMPETITORS = {
    "Eneba":"https://www.eneba.com/",
    "DoctorSIM":"https://www.doctorsim.com/",
    "KarteDirekt":"https://kartedirekt.de/","Aufladen":"https://aufladen.de/",
}

# Blog URL overrides for competitors whose blog is on a different subdomain or non-standard path
BLOG_OVERRIDES = {
    "DoctorSIM": ["https://blog-en.doctorsim.com/"],
    "Eneba": ["https://www.eneba.com/hub/"],
    "KarteDirekt": ["https://kartedirekt.de/blogs"],
}

WIKI_PAGES = [
    "Grand_Theft_Auto_VI","EA_Sports_FC","Fortnite","PlayStation_5",
    "Xbox_Game_Pass","Nintendo_Switch_2","Genshin_Impact","Call_of_Duty",
    "Minecraft","Roblox","Valorant","League_of_Legends","Steam_(service)",
    "Spotify","Netflix","Crunchyroll","Discord","Honkai:_Star_Rail",
    "Elden_Ring","Monster_Hunter_Wilds",
]

TREND_BATCHES = [
    ["EA FC 26","Fortnite","GTA 6","PS Plus","Game Pass"],
    ["Steam sale","Nintendo Switch 2","Roblox","Genshin Impact","Valorant"],
    ["Spotify Premium","Netflix","Discord Nitro","gift card","V-Bucks"],
    ["Call of Duty","Minecraft","Xbox","PlayStation","Steam Deck"],
    ["Crunchyroll","anime 2026","Monster Hunter Wilds","Elden Ring","Destiny 2"],
]

# Auto-generate news topics from every KW category + extras
_AUTO_TOPICS = []
for _cat, _kws in KW.items():
    _AUTO_TOPICS.append(f"{_cat} update")
    _AUTO_TOPICS.append(f"{_cat} news")
    if _kws: _AUTO_TOPICS.append(f"{_kws[0]} event")

NEWS_TOPICS = list(dict.fromkeys(_AUTO_TOPICS + [
    "EA FC 26 TOTY","EA FC 26 promo","EA FC 26 Future Stars",
    f"PlayStation Plus {MON} {YEAR}","PS Plus free games",
    "PlayStation State of Play","PS5 Pro",
    f"Xbox Game Pass {MON} {YEAR}","Xbox Developer Direct",
    "Nintendo Direct","Nintendo Switch 2",
    "Fortnite new season","GTA 6","GTA 6 release date",
    "Call of Duty new season","Warzone new season",
    "Genshin Impact new banner","Honkai Star Rail update",
    "Overwatch 2 event","Overwatch 2 season",
    "Diablo 4 season","Hearthstone expansion",
    "gaming gift card deals","game release date",
    "esports tournament","Epic Games free",
    "anime new season","One Piece","Dragon Ball",
    "Monster Hunter Wilds","Elden Ring DLC",
    "GOG sale","Humble Bundle","free game giveaway",
    "new game release this week","gaming news today",
    "biggest game releases this month","trending games",
    "game awards","gaming event this week",
]))

# =============================================================================
# SECTION 1 - DEPENDENCIES
# =============================================================================

def _install():
    pkgs = ["google-genai","pytrends","python-docx","gspread","google-auth",
            "requests","beautifulsoup4","feedparser","cryptography"]
    for p in pkgs:
        mod = p.replace("-","_").split("[")[0]
        if mod == "python_docx": mod = "docx"
        if mod == "google_genai": mod = "google.genai"
        try: __import__(mod)
        except ImportError:
            subprocess.check_call([sys.executable,"-m","pip","install",p,"-q"])

print("Installing deps...", end=" "); _install(); print("OK")
try:
    from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    HAS_CRYPTO = True
except ImportError: HAS_CRYPTO = False

import requests
from bs4 import BeautifulSoup
import feedparser
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from google import genai
from google.genai import types as gtypes

GCLIENT = genai.Client(api_key=GEMINI_KEY)
GEMINI_MODEL = "gemini-3.1-pro-preview"

try:
    import gspread; from google.auth import default as _gauth
    HAS_GSPREAD = True
except ImportError: HAS_GSPREAD = False

try:
    from pytrends.request import TrendReq
    HAS_PYTRENDS = True
except ImportError: HAS_PYTRENDS = False

# =============================================================================
# SECTION 2 - DATA MODELS
# =============================================================================

@dataclass
class Signal:
    source: str; title: str; desc: str = ""; url: str = ""
    score: float = 0.0; meta: dict = field(default_factory=dict)

@dataclass
class Candidate:
    title: str; signals: List[Signal] = field(default_factory=list)
    sources: int = 0; score: float = 0.0; url: str = ""
    category: str = ""; categories: List[str] = field(default_factory=list)
    biz_category: str = "GMG"; biz_categories: List[str] = field(default_factory=list)

@dataclass
class ScoreSpec:
    """How to normalize raw scores from a source to 0-100."""
    floor: float = 0.0
    ceiling: float = 100.0

SCORE_SPECS = {
    "trends": ScoreSpec(0, 100), "reddit": ScoreSpec(50, 80),
    "steam": ScoreSpec(40, 100), "wiki": ScoreSpec(0, 100),
    "youtube": ScoreSpec(45, 75), "news": ScoreSpec(40, 80),
    "oxylabs_news": ScoreSpec(70, 95), "competitor": ScoreSpec(35, 60),
    "cheapshark": ScoreSpec(35, 100), "steamspy": ScoreSpec(0, 100),
    "gamerpower": ScoreSpec(50, 80), "epic": ScoreSpec(45, 80),
    "gog": ScoreSpec(40, 85), "humble": ScoreSpec(45, 85),
    "freetogame": ScoreSpec(30, 60), "anime": ScoreSpec(0, 100),
    "sitemap": ScoreSpec(30, 55),
}

def normalize_scores(all_sig):
    """Normalize all signal scores to a consistent 0-100 scale per source."""
    for source, signals in all_sig.items():
        if not signals: continue
        spec = SCORE_SPECS.get(source, ScoreSpec())
        raw_scores = [s.score for s in signals]
        src_min, src_max = min(raw_scores), max(raw_scores)
        for s in signals:
            if src_max > src_min:
                normalized = spec.floor + (s.score - src_min) / (src_max - src_min) * (spec.ceiling - spec.floor)
            else:
                normalized = (spec.floor + spec.ceiling) / 2
            s.score = round(min(100, max(0, normalized)), 1)

# =============================================================================
# SECTION 3 - UTILITIES
# =============================================================================

import threading as _threading

class RateLimiter:
    """Simple token-bucket rate limiter for API calls."""
    def __init__(self, calls_per_minute=10):
        self._interval = 60.0 / calls_per_minute
        self._last_call = 0.0
        self._lock = _threading.Lock()
    def wait(self):
        with self._lock:
            now = time.time()
            elapsed = now - self._last_call
            if elapsed < self._interval:
                time.sleep(self._interval - elapsed)
            self._last_call = time.time()

GEMINI_LIMITER = RateLimiter(calls_per_minute=8)

def GET(url, headers=None, timeout=15, retries=2):
    h = headers or {"User-Agent": UA}
    for i in range(retries+1):
        try:
            r = requests.get(url, headers=h, timeout=timeout)
            if r.status_code == 200: return r
            if r.status_code == 429:
                retry_after = r.headers.get("Retry-After")
                if retry_after:
                    try: delay = min(float(retry_after), 30)
                    except ValueError: delay = min(1.5 * (2 ** i), 10)
                else: delay = min(1.5 * (2 ** i), 10)
                log.info(f"Rate limited (429) on {url[:60]}, waiting {delay:.1f}s")
                time.sleep(delay); continue
            if r.status_code in (500, 502, 503, 504):
                delay = min(1.5 * (2 ** i), 10)
                log.debug(f"Server error {r.status_code} on {url[:60]}, retry {i+1}/{retries}")
                time.sleep(delay); continue
            log.debug(f"GET {url[:60]} returned {r.status_code}")
            return None
        except requests.exceptions.Timeout:
            log.debug(f"Timeout on {url[:60]} (attempt {i+1}/{retries+1})")
            if i < retries: time.sleep(1.5 * (2 ** i))
        except requests.exceptions.ConnectionError as e:
            log.debug(f"Connection error on {url[:60]}: {e}")
            if i < retries: time.sleep(1.5 * (2 ** i))
        except Exception as e:
            log.warning(f"Unexpected error on GET {url[:60]}: {type(e).__name__}: {e}")
            if i < retries: time.sleep(1.5 * (2 ** i))
    return None

def fuzz(a, b, t=FUZZ_T):
    a2 = re.sub(r'[^a-z0-9 ]','',a.lower())
    b2 = re.sub(r'[^a-z0-9 ]','',b.lower())
    if not a2 or not b2: return False
    # Number guard: if both contain numbers and they differ, reject
    nums_a = re.findall(r'\d+', a2)
    nums_b = re.findall(r'\d+', b2)
    if nums_a and nums_b and nums_a != nums_b: return False
    return SequenceMatcher(None,a2,b2).ratio() >= t

def cats(text):
    lo = text.lower()
    matched = []
    for c, kws in KW.items():
        hits = [k for k in kws if k.lower() in lo]
        if hits:
            matched.append(c)
            log.debug(f"  KW match: '{text[:60]}' -> {c} (matched: {', '.join(hits[:3])})")
    return matched

def biz_cats(text):
    """Return list of business categories (GMG/ENT/PPM/MTU) for given text."""
    kw_cats = cats(text)
    biz = list(dict.fromkeys(BIZ_CATS.get(c, "GMG") for c in kw_cats))
    return biz if biz else ["GMG"]

def recent(ds, days=7):
    if not ds: return False
    cut = NOW - timedelta(days=days)
    for f in ['%a, %d %b %Y %H:%M:%S %z','%a, %d %b %Y %H:%M:%S %Z',
              '%Y-%m-%dT%H:%M:%S%z','%Y-%m-%dT%H:%M:%SZ',
              '%Y-%m-%dT%H:%M:%S.%f%z','%Y-%m-%d %H:%M:%S','%Y-%m-%d']:
        try:
            p = datetime.strptime(ds.replace('GMT','+0000'),f)
            if p.tzinfo: p = p.replace(tzinfo=None)
            return p >= cut
        except (ValueError, TypeError): continue
    return False

def mass_appeal(t):
    lo = t.lower()
    return not any(x in lo for x in
        ["my ","i ","i'm","i've","me ","just got","finally",
         "my friend","helped me","unpopular opinion","does anyone",
         "question about","need help","eli5","ama"])

def norm(v, mx): return min(100.0, v/mx*100.0) if mx > 0 else 0.0
def esc(s): return _html.escape(str(s))

def _best_url(signals):
    prio = {"news":1,"oxylabs_news":1,"reddit":2,"youtube":3,"steam":4,"gog":5,"humble":6,
            "cheapshark":7,"epic":8,"anime":9,"gamerpower":10,"steamspy":11,
            "freetogame":12,"competitor":13,"trends":14,"wiki":99}
    urls = [(s.url, prio.get(s.source,20)) for s in signals if s.url]
    if not urls: return ""
    urls.sort(key=lambda x: x[1])
    return urls[0][0]

def _fmt_action(a):
    if isinstance(a, dict):
        owner = a.get("owner","")
        action = a.get("action","")
        due = a.get("due","")
        parts = []
        if owner: parts.append(f"<strong>{esc(owner)}</strong>: ")
        parts.append(esc(action))
        if due: parts.append(f" <em>({esc(due)})</em>")
        return "".join(parts)
    return esc(str(a))

def _match_cand(title, cands):
    """Find BEST matching candidate (not first-above-threshold). Prevents wrong URL matches."""
    best = None; best_r = 0
    a = re.sub(r'[^a-z0-9 ]','',title.lower())
    if not a: return None
    for c in cands:
        b = re.sub(r'[^a-z0-9 ]','',c.title.lower())
        if not b: continue
        r = SequenceMatcher(None,a,b).ratio()
        if r > best_r: best_r = r; best = c
    return best if best_r > 0.45 else None

def _domain(url):
    if not url: return ""
    try: return urlparse(url).netloc.replace("www.","")[:30]
    except Exception: return ""

# =============================================================================
# SECTION 4 - 16 DATA FETCHERS
# =============================================================================

class TrendsFetcher:
    def fetch(self):
        if not HAS_PYTRENDS:
            log.info("I wanted to check Google Trends but pytrends isn't installed. Skipping."); return []
        out = []
        log.info(f"I'm checking Google Trends for search interest across {len(TREND_BATCHES)} keyword batches...")
        try:
            pt = TrendReq(hl='en-US',tz=360,retries=2,backoff_factor=0.5,
                          requests_args={'headers':{'Cookie':'CONSENT=YES+'}})
            try:
                tr = pt.trending_searches(pn='united_states')
                for _,row in tr.head(20).iterrows():
                    q = str(row[0]); c = cats(q)
                    if c: out.append(Signal("trends",q,"Trending search (US)",score=80,meta={"cats":c}))
            except Exception as e: log.debug(f"Trends top searches: {e}")
            for batch in TREND_BATCHES:
                try:
                    pt.build_payload(batch,cat=0,timeframe='now 7-d',geo='',gprop='')
                    df = pt.interest_over_time()
                    if df.empty: continue
                    for kw in batch:
                        if kw not in df.columns: continue
                        vals = df[kw].tolist(); cur = float(vals[-1]) if vals else 0
                        out.append(Signal("trends",kw,f"Search interest: {cur:.0f}/100",
                            score=cur,meta={"cats":cats(kw)}))
                    time.sleep(2)
                except Exception as e: log.debug(f"Trends batch: {e}"); time.sleep(3)
        except Exception as e: log.debug(f"Trends init: {e}")
        if out:
            top = sorted(out, key=lambda s: s.score, reverse=True)[:3]
            log.info(f"Found {len(out)} trend signals. Hottest: {', '.join(s.title for s in top)}")
        else:
            log.info("No trend data this time — Google may be rate-limiting us.")
        return out

class RedditFetcher:
    def fetch(self):
        out = []
        log.info(f"I'm browsing {len(SUBREDDITS)} gaming subreddits to see what people are discussing...")
        subs_with_results = set()
        for sub in SUBREDDITS:
            try:
                feed = feedparser.parse(f"https://www.reddit.com/r/{sub}/hot/.rss?limit=8",
                    request_headers={"User-Agent":"RechargeScanner/4.2"})
                for e in feed.entries[:5]:
                    t = e.get("title","")
                    if not mass_appeal(t): continue
                    cc = cats(t)
                    if not cc: continue
                    out.append(Signal("reddit",t[:150],f"r/{sub} (Hot)",
                        url=e.get("link",""),score=65,meta={"sub":sub,"cats":cc}))
                    subs_with_results.add(sub)
                time.sleep(2)
            except Exception as e:
                log.warning(f"Reddit r/{sub}: {e}"); continue
        top = sorted(out, key=lambda s: s.score, reverse=True)[:3]
        top_titles = ', '.join(s.title[:50] for s in top) if top else "nothing notable"
        log.info(f"Found {len(out)} posts across {len(subs_with_results)} subreddits. Hot topics: {top_titles}")
        return out

class SteamFetcher:
    """Merged: featured + top_sellers + specials + new_releases + coming_soon."""
    def fetch(self):
        out = []
        log.info("I'm checking Steam's store for trending games, sales, and new releases...")
        try:
            r = GET("https://store.steampowered.com/api/featured/")
            if r:
                for it in r.json().get("featured_win",[])[:10]:
                    n,d,aid = it.get("name",""),it.get("discount_percent",0),it.get("id","")
                    out.append(Signal("steam",n,f"Featured{f', {d}% off' if d else ''}",
                        url=f"https://store.steampowered.com/app/{aid}",
                        score=min(60+d*0.4,100),meta={"discount":d,"cats":cats(n)}))
        except Exception as e: log.debug(f"Steam featured: {e}")
        try:
            r = GET("https://store.steampowered.com/api/featuredcategories/")
            if r:
                data = r.json()
                sections = {"top_sellers":70,"specials":50,"new_releases":60,"coming_soon":45}
                for sec,base in sections.items():
                    for it in data.get(sec,{}).get("items",[])[:10]:
                        n,d,aid = it.get("name",""),it.get("discount_percent",0),it.get("id","")
                        label = sec.replace('_',' ').title()
                        out.append(Signal("steam",n,f"{label}{f', {d}% off' if d else ''}",
                            url=f"https://store.steampowered.com/app/{aid}",
                            score=min(base+d*0.3,100),
                            meta={"discount":d,"section":sec,"cats":cats(n) or ["Steam"]}))
        except Exception as e: log.debug(f"Steam categories: {e}")
        top = sorted(out, key=lambda s: s.score, reverse=True)[:3]
        top_titles = ', '.join(s.title[:40] for s in top) if top else "no results"
        log.info(f"Found {len(out)} games on Steam. Top: {top_titles}")
        return out

class WikiFetcher:
    def fetch(self):
        out = []
        log.info(f"I'm checking Wikipedia page views for {len(WIKI_PAGES)} key topics...")
        end = NOW-timedelta(days=1); s1 = end-timedelta(days=6)
        fmt = lambda d: d.strftime("%Y%m%d")
        hd = {"User-Agent":"RechargeScanner/4.2 (content-research)"}
        for pg in WIKI_PAGES:
            try:
                r = GET(f"https://wikimedia.org/api/rest_v1/metrics/pageviews/per-article/en.wikipedia/all-access/all-agents/{pg}/daily/{fmt(s1)}/{fmt(end)}",headers=hd,timeout=10)
                if r:
                    views = sum(d.get("views",0) for d in r.json().get("items",[]))
                    name = pg.replace("_"," ")
                    if views > 1000:
                        out.append(Signal("wiki",name,f"{views:,} views this week",
                            url=f"https://en.wikipedia.org/wiki/{pg}",
                            score=norm(views,200000),meta={"views":views,"cats":cats(name)}))
            except Exception as e:
                log.warning(f"Wiki {pg}: {e}"); continue
        top = sorted(out, key=lambda s: s.meta.get("views",0), reverse=True)[:3]
        top_str = ', '.join(f"{s.title} ({s.meta.get('views',0):,} views)" for s in top) if top else "none"
        log.info(f"Most viewed: {top_str}")
        return out

class YTFetcher:
    def fetch(self):
        out = []; ns = {"a":"http://www.w3.org/2005/Atom"}
        log.info(f"I'm checking {len(YT_CHANNELS)} YouTube gaming channels for new videos...")
        for ch,cid in YT_CHANNELS.items():
            try:
                r = GET(f"https://www.youtube.com/feeds/videos.xml?channel_id={cid}",timeout=10)
                if not r: continue
                root = ET.fromstring(r.content)
                for e in root.findall("a:entry",ns)[:10]:
                    te = e.find("a:title",ns); le = e.find("a:link",ns); pe = e.find("a:published",ns)
                    if te is None: continue
                    t = te.text or ""; u = le.get("href","") if le is not None else ""
                    pub = pe.text if pe is not None else ""
                    if not recent(pub,7): continue
                    cc = cats(t)
                    if cc: out.append(Signal("youtube",t[:150],f"YouTube: {ch}",url=u,score=60,meta={"ch":ch,"cats":cc}))
            except Exception as e:
                log.warning(f"YouTube {ch}: {e}"); continue
        top = out[:3]
        top_titles = ', '.join(s.title[:45] for s in top) if top else "nothing recent"
        log.info(f"Found {len(out)} videos. Latest: {top_titles}")
        return out

class NewsFetcher:
    def fetch(self):
        out = []; seen = set()
        log.info(f"I'm scanning {len(RSS_FEEDS)} news sites (IGN, GameSpot, PYMNTS, etc.) for headlines...")
        # Skip Google News RSS if Oxylabs handles it (avoids massive overlap)
        skip_gnews = bool(OXYLABS_USER and OXYLABS_PASS)
        for topic in ([] if skip_gnews else NEWS_TOPICS):
            try:
                feed = feedparser.parse(f"https://news.google.com/rss/search?q={quote(topic)}+when:7d&hl=en-US&gl=US&ceid=US:en")
                for e in feed.entries[:3]:
                    t = e.get("title",""); src = "News"
                    if " - " in t: t,src = t.rsplit(" - ",1)
                    k = re.sub(r'[^a-z0-9]','',t[:80].lower())
                    if k in seen: continue
                    seen.add(k)
                    if not mass_appeal(t) or not recent(e.get("published",""),7): continue
                    cc = cats(t)
                    if cc: out.append(Signal("news",t[:150],f"via {src}",url=e.get("link",""),score=70,meta={"src":src,"cats":cc}))
                time.sleep(0.15)
            except Exception as e:
                log.debug(f"News topic '{topic[:30]}': {e}"); continue
        for fn,fu in RSS_FEEDS.items():
            try:
                feed = feedparser.parse(fu)
                for e in feed.entries[:20]:
                    t = e.get("title","")
                    k = re.sub(r'[^a-z0-9]','',t[:80].lower())
                    if k in seen: continue
                    seen.add(k)
                    if not mass_appeal(t) or not recent(e.get("published",e.get("updated","")),7): continue
                    cc = cats(t)
                    if cc:
                        out.append(Signal("news",t[:150],f"via {fn}",url=e.get("link",""),score=65,meta={"src":fn,"cats":cc}))
                    elif fn in ("IGN","GameSpot","Kotaku","PC Gamer","Eurogamer","Polygon","GamesRadar","Dexerto","VG247","DualShockers","GameRant","GamesIndustry.biz","Screen Rant","PYMNTS","What's On Netflix","VGC","PCGamesN","Collider","Deadline TV","CinemaBlend","ComingSoon","Digital Spy"):
                        out.append(Signal("news",t[:150],f"via {fn}",url=e.get("link",""),score=45,meta={"src":fn,"cats":["General"]}))
                time.sleep(0.05)
            except Exception as e:
                log.debug(f"News RSS {fn}: {e}"); continue
        top = sorted(out, key=lambda s: s.score, reverse=True)[:3]
        top_titles = ', '.join(s.title[:45] for s in top) if top else "none"
        log.info(f"Collected {len(out)} articles. Headlines: {top_titles}")
        return out

class OxylabsNewsFetcher:
    """Fresh real-time news via Oxylabs Web Scraper API (Google News)."""
    # Focused queries per business category for maximum fresh coverage
    QUERIES = {
        "GMG": [
            "gaming news today", "PlayStation news", "Xbox Game Pass news",
            "Nintendo Switch 2", "Steam sale news", "Fortnite update",
            "GTA 6 news", "Call of Duty news", "EA FC news",
            "Genshin Impact update", "Roblox news", "Valorant update",
            "gaming gift card deals", "esports news today",
            "Razer Gold news", "PlayStation Store deals",
        ],
        "ENT": [
            "Spotify news today", "Netflix new releases", "Disney Plus news",
            "Crunchyroll anime news", "streaming service news",
            "YouTube Premium news", "Amazon Prime Video news",
            "TikTok news today", "Apple gift card news",
        ],
        "PPM": [
            "gift card deals today", "Paysafecard news",
            "digital payment news", "prepaid card news",
            "Transcash news", "Neosurf news", "Flexepin news",
            "Jeton Cash news", "CashLib prepaid news",
            "Google Play gift card news", "Amazon gift card deals",
            "Visa gift card news",
        ],
        "MTU": [
            "mobile top up news", "phone credit deals",
            "mobile carrier news", "prepaid mobile news",
            "Lycamobile news", "Digimobil news",
            "Lebara news", "T-Mobile prepaid news",
        ],
    }

    # Top Recharge.com product pages (from Search Console data) — used for targeted news
    PRODUCT_QUERIES = [
        ("Paysafecard", "PPM"), ("Roblox gift card", "GMG"), ("Steam gift card", "GMG"),
        ("Transcash recharge", "PPM"), ("Lycamobile recharge", "MTU"),
        ("Neosurf voucher", "PPM"), ("Google Play gift card", "PPM"),
        ("PUBG Mobile UC", "GMG"), ("TikTok coins", "ENT"),
        ("Spotify Premium gift card", "ENT"), ("Netflix gift card", "ENT"),
        ("PlayStation Store card", "GMG"), ("Xbox gift card", "GMG"),
        ("Fortnite V-Bucks", "GMG"), ("Apple gift card", "ENT"),
        ("Amazon gift card", "PPM"), ("Razer Gold", "GMG"),
    ]

    def fetch(self):
        if not OXYLABS_USER or not OXYLABS_PASS:
            log.info("I wanted to search Google News via Oxylabs but credentials aren't set. Skipping."); return []
        out = []; seen = set()
        # Flatten and deduplicate queries, prioritize GMG/ENT (more results expected)
        all_queries = []
        for bc in ["GMG","ENT","PPM","MTU"]:
            for q in self.QUERIES.get(bc, []):
                all_queries.append((q, bc))
        # Add product-specific queries (from Recharge.com top products)
        for pq, pbc in self.PRODUCT_QUERIES:
            all_queries.append((pq + " news", pbc))
        log.info(f"Now searching Google News for {len(all_queries)} queries across Gaming, Entertainment, Payments, Mobile...")
        for idx, (query, bc) in enumerate(all_queries, 1):
            if idx % 15 == 0:
                log.info(f"  ...{idx}/{len(all_queries)} done, found {len(out)} articles so far")
            try:
                payload = {
                    "source": "google_search",
                    "query": query,
                    "parse": True,
                    "context": [
                        {"key": "tbm", "value": "nws"},
                        {"key": "tbs", "value": "qdr:d"},  # last 24 hours
                    ],
                    "geo_location": "United States",
                    "locale": "en-US",
                    "limit": 10,
                }
                r = requests.post("https://realtime.oxylabs.io/v1/queries",
                    auth=(OXYLABS_USER, OXYLABS_PASS), json=payload, timeout=30)
                if r.status_code != 200: continue
                data = r.json()
                results = data.get("results", [])
                if not results: continue
                content = results[0].get("content", {})
                if not isinstance(content, dict): continue
                main_items = content.get("results", {}).get("main", [])
                for item in main_items[:8]:
                    title = item.get("title", "").strip()
                    if not title: continue
                    k = re.sub(r'[^a-z0-9]','', title[:80].lower())
                    if k in seen: continue
                    seen.add(k)
                    if not mass_appeal(title): continue
                    url = item.get("url", "")
                    source = item.get("source", "Google News")
                    age = item.get("relative_publish_date", "")
                    cc = cats(title)
                    # Higher score for very fresh news (hours ago)
                    score = 80  # base: higher than RSS news (70)
                    if "minute" in age.lower(): score = 90
                    elif "hour" in age.lower():
                        try:
                            hrs = int(re.search(r'(\d+)', age).group(1))
                            score = 90 if hrs <= 3 else 85
                        except: score = 85
                    out.append(Signal("oxylabs_news", title[:150],
                        f"via {source} ({age})", url=url, score=score,
                        meta={"src": source, "age": age, "cats": cc if cc else [query.split()[0]],
                              "biz_cat": bc, "fresh": True}))
                time.sleep(0.3)  # rate limit
            except Exception as e:
                log.warning(f"OxylabsNews query '{query}': {e}"); continue
        log.info(f"Google News search done. Found {len(out)} fresh articles."); return out

class CompetitorFetcher:
    DEAL_PATHS = ["/deals","/promotions","/sale","/hot-deals","/best-deals"]

    def fetch(self):
        out = []; hd = {"User-Agent":UA,"Accept":"text/html","Accept-Language":"en-US,en;q=0.9"}
        log.info(f"I'm visiting competitor sites ({', '.join(COMPETITORS.keys())}) to see promotions...")
        for name,url in COMPETITORS.items():
            try:
                # Homepage scan
                r = GET(url,headers=hd,timeout=12)
                if not r: continue
                soup = BeautifulSoup(r.text,"html.parser")
                text = soup.get_text(" ",strip=True)[:5000]; promo = []
                for cat,kws in KW.items():
                    if any(k.lower() in text.lower() for k in kws) and cat not in promo: promo.append(cat)
                # Extract headline promotions
                hero_texts = []
                for tag in soup.find_all(["h1","h2","h3","title"]):
                    tag_text = tag.get_text(strip=True)
                    if tag_text and len(tag_text) > 5:
                        hero_texts.append(tag_text[:100])
                    for c in cats(tag_text):
                        if c not in promo: promo.append(c)
                # Detect sale percentages
                sale_matches = re.findall(r'(\d{1,3})%\s*(?:off|discount|sale)', text[:3000], re.I)
                sale_note = f" (up to {max(int(x) for x in sale_matches)}% off)" if sale_matches else ""
                for p in promo:
                    bc = BIZ_CATS.get(p, "GMG")
                    out.append(Signal("competitor",f"{name}: {p}{sale_note}",
                        f"{name} promoting {p}",url=url,score=45,
                        meta={"comp":name,"product":p,"cats":[p],"biz_cat":bc,
                              "activity_type":"promotion","hero":hero_texts[:3]}))
                # Try deals/promotions pages
                domain = urlparse(url).scheme + "://" + urlparse(url).netloc
                for path in self.DEAL_PATHS:
                    try:
                        r2 = GET(domain + path, headers=hd, timeout=8)
                        if not r2 or r2.status_code != 200: continue
                        soup2 = BeautifulSoup(r2.text, "html.parser")
                        for tag in soup2.find_all(["h1","h2","h3"])[:10]:
                            tag_text = tag.get_text(strip=True)
                            cc = cats(tag_text)
                            if cc:
                                for c in cc:
                                    if c not in promo: promo.append(c)
                                out.append(Signal("competitor",f"{name} deals: {tag_text[:80]}",
                                    f"{name} deals page",url=domain+path,score=50,
                                    meta={"comp":name,"product":cc[0],"cats":cc,
                                          "biz_cat":BIZ_CATS.get(cc[0],"GMG"),"activity_type":"deal"}))
                        break  # found working deals page
                    except Exception as e:
                        log.debug(f"Competitor {name} deals {path}: {e}"); continue
                time.sleep(1)
            except Exception as e:
                log.warning(f"Competitor {name}: {e}"); continue
        log.info(f"Found {len(out)} promotional signals from competitors."); return out

class CheapSharkFetcher:
    def fetch(self):
        out = []
        log.info("I'm checking CheapShark for PC game deals...")
        try:
            r = GET("https://www.cheapshark.com/api/1.0/deals?storeID=1&upperPrice=15&pageSize=30&sortBy=Deal+Rating",timeout=10)
            if r:
                for d in r.json()[:30]:
                    name,savings = d.get("title",""),float(d.get("savings",0))
                    normal,sale = float(d.get("normalPrice",0)),float(d.get("salePrice",0))
                    cc = cats(name); sc = min(40+savings*0.6,100)
                    out.append(Signal("cheapshark",name,f"${sale:.0f} (was ${normal:.0f}, {savings:.0f}% off)",
                        url=f"https://www.cheapshark.com/redirect?dealID={d.get('dealID','')}",
                        score=sc,meta={"savings":savings,"cats":cc if cc else ["Gaming"]}))
        except Exception as e: log.warning(f"CheapShark: {e}")
        top = sorted(out, key=lambda s: s.score, reverse=True)[:3]
        top_str = ', '.join(f"{s.title[:30]} ({s.desc})" for s in top) if top else "none"
        log.info(f"Found {len(out)} deals. Best: {top_str}")
        return out

class SteamSpyFetcher:
    def fetch(self):
        out = []
        log.info("I'm checking SteamSpy for the most-played games this week...")
        try:
            r = GET("https://steamspy.com/api.php?request=top100in2weeks",timeout=12)
            if r:
                for appid,info in sorted(r.json().items(),key=lambda x:x[1].get("ccu",0),reverse=True)[:20]:
                    name,ccu,players = info.get("name",""),info.get("ccu",0),info.get("players_2weeks",0)
                    cc = cats(name); sc = norm(ccu,500000)
                    out.append(Signal("steamspy",name,f"{ccu:,} playing now, {players:,} in 2wk",
                        url=f"https://store.steampowered.com/app/{appid}",
                        score=sc,meta={"ccu":ccu,"cats":cc if cc else ["Steam"]}))
        except Exception as e: log.warning(f"SteamSpy: {e}")
        top_game = max(out, key=lambda s: s.meta.get("ccu",0)).title if out else "none"
        log.info(f"Found {len(out)} popular games. Most played: {top_game}")
        return out

class GamerPowerFetcher:
    def fetch(self):
        out = []
        log.info("I'm checking GamerPower for free game giveaways...")
        try:
            r = GET("https://www.gamerpower.com/api/giveaways?sort-by=popularity",timeout=10)
            if r:
                for g in r.json()[:25]:
                    name,platforms,worth = g.get("title",""),g.get("platforms",""),g.get("worth","N/A")
                    cc = cats(name+" "+platforms)
                    out.append(Signal("gamerpower",name,f"{g.get('type','')} on {platforms} ({worth})",
                        url=g.get("open_giveaway_url",""),score=65,
                        meta={"platforms":platforms,"cats":cc if cc else ["Gaming"]}))
        except Exception as e: log.warning(f"GamerPower: {e}")
        log.info(f"Found {len(out)} active giveaways.")
        return out

class EpicFreeFetcher:
    def fetch(self):
        out = []
        log.info("I'm checking Epic Games Store for this week's free games...")
        try:
            r = GET("https://store-site-backend-static-ipv4.ak.epicgames.com/freeGamesPromotions?locale=en-US&country=US&allowCountries=US",timeout=10)
            if r:
                for el in r.json().get("data",{}).get("Catalog",{}).get("searchStore",{}).get("elements",[]):
                    name = el.get("title",""); promos = el.get("promotions")
                    if not promos: continue
                    if promos.get("promotionalOffers"):
                        out.append(Signal("epic",name,"FREE NOW on Epic",url="https://store.epicgames.com/en-US/free-games",
                            score=75,meta={"status":"free_now","cats":cats(name) or ["Gaming"]}))
                    elif promos.get("upcomingPromotionalOffers"):
                        out.append(Signal("epic",name,"Coming free on Epic",url="https://store.epicgames.com/en-US/free-games",
                            score=55,meta={"status":"upcoming","cats":cats(name) or ["Gaming"]}))
        except Exception as e: log.warning(f"Epic: {e}")
        free_names = [s.title for s in out if "FREE NOW" in s.desc]
        if free_names:
            log.info(f"Epic has {len(free_names)} free games: {', '.join(free_names)}")
        else:
            log.info(f"Found {len(out)} Epic games ({len([s for s in out if 'Coming' in s.desc])} upcoming free).")
        return out

class GOGFetcher:
    def fetch(self):
        out = []
        log.info("I'm checking GOG for popular games and sales...")
        try:
            r = GET("https://www.gog.com/games/ajax/filtered?mediaType=game&page=1&sort=popularity&limit=20",
                headers={"User-Agent":UA,"Accept":"application/json"},timeout=12)
            if r:
                for g in r.json().get("products",[])[:20]:
                    name = g.get("title",""); price = g.get("price",{})
                    discount = price.get("discount",0) if isinstance(price,dict) else 0
                    slug = g.get("url","")
                    gog_url = f"https://www.gog.com{slug}" if slug else "https://www.gog.com"
                    sc = 50+(discount if isinstance(discount,int) else 0)*0.3
                    out.append(Signal("gog",name,f"GOG Popular{f', {discount}% off' if discount else ''}",
                        url=gog_url,score=min(sc,100),meta={"cats":cats(name) or ["Gaming"]}))
        except Exception as e: log.warning(f"GOG: {e}")
        log.info(f"Found {len(out)} games on GOG.")
        return out

class HumbleFetcher:
    def fetch(self):
        out = []
        log.info("I'm checking Humble Bundle for bestsellers...")
        try:
            r = GET("https://www.humblebundle.com/store/api/search?sort=bestselling&filter=all&hmb_source=store_navbar&page=0",
                headers={"User-Agent":UA,"Accept":"application/json"},timeout=12)
            if r:
                for g in r.json().get("results",[])[:20]:
                    name = g.get("human_name","") or g.get("human_url","")
                    slug = g.get("human_url","")
                    cp = (g.get("current_price") or {}).get("amount",0)
                    fp = (g.get("full_price") or {}).get("amount",0)
                    discount = int((1-cp/fp)*100) if fp>0 else 0
                    url = f"https://www.humblebundle.com/store/{slug}" if slug else "https://www.humblebundle.com/store"
                    out.append(Signal("humble",name,f"Humble bestseller{f', {discount}% off' if discount>0 else ''}",
                        url=url,score=min(55+discount*0.3,100),meta={"cats":cats(name) or ["Gaming"]}))
        except Exception as e: log.warning(f"Humble: {e}")
        log.info(f"Found {len(out)} Humble bestsellers.")
        return out

class FreeToGameFetcher:
    def fetch(self):
        out = []
        log.info("I'm checking the free-to-play game directory...")
        try:
            r = GET("https://www.freetogame.com/api/games?sort-by=relevance",timeout=10)
            if r:
                for g in r.json()[:20]:
                    name,genre,platform = g.get("title",""),g.get("genre",""),g.get("platform","")
                    out.append(Signal("freetogame",name,f"Free {genre} on {platform}",
                        url=g.get("game_url",""),score=45,
                        meta={"cats":cats(name+" "+genre) or ["Gaming"]}))
        except Exception as e: log.warning(f"FreeToGame: {e}")
        log.info(f"Found {len(out)} free-to-play games.")
        return out

class AnimeFetcher:
    def fetch(self):
        out = []
        log.info("I'm checking top airing and upcoming anime (for Crunchyroll insights)...")
        try:
            r = GET("https://api.jikan.moe/v4/top/anime?filter=airing&limit=15",timeout=12)
            if r:
                for a in r.json().get("data",[]):
                    name,score,members = a.get("title",""),a.get("score",0) or 0,a.get("members",0)
                    out.append(Signal("anime",name,f"Top airing, MAL {score}, {members:,} fans",
                        url=a.get("url",""),score=norm(members,1000000),
                        meta={"cats":cats(name) or ["Crunchyroll"]}))
        except Exception as e: log.debug(f"Anime airing: {e}")
        try:
            time.sleep(1.5)
            r = GET("https://api.jikan.moe/v4/seasons/upcoming?limit=10",timeout=12)
            if r:
                for a in r.json().get("data",[]):
                    name,members = a.get("title",""),a.get("members",0)
                    out.append(Signal("anime",name,f"Upcoming, {members:,} anticipating",
                        url=a.get("url",""),score=norm(members,500000),
                        meta={"cats":cats(name) or ["Crunchyroll"]}))
        except Exception as e: log.debug(f"Anime upcoming: {e}")
        top = sorted(out, key=lambda s: s.score, reverse=True)[:3]
        top_titles = ', '.join(s.title[:35] for s in top) if top else "none"
        log.info(f"Found {len(out)} anime series. Top: {top_titles}")
        return out

class SitemapFetcher:
    """Fetch competitor sitemaps + blog RSS to find pages published this week."""
    BLOG_PATHS = [
        "/blog/feed","/blog/rss","/feed","/rss","/feed.xml","/rss.xml",
        "/blog/feed/","/blog/atom.xml","/en/blog/feed","/news/feed",
        "/blog.rss","/articles/feed","/feed/rss2","/blog/rss.xml",
        "/hub/feed","/hub/rss","/hub/feed.xml",
    ]

    def fetch(self):
        out = []; cutoff = NOW - timedelta(days=14)
        blog_names = ', '.join(BLOG_OVERRIDES.keys()) if BLOG_OVERRIDES else 'competitors'
        log.info(f"I'm reading competitor blogs ({blog_names})...")
        for name, base_url in SITEMAP_COMPETITORS.items():
            try:
                domain = urlparse(base_url).netloc
                scheme = "https"
                log.debug(f"  --- {name} ({domain}) ---")

                # === PART A: Blog RSS feeds ===
                # Check BLOG_OVERRIDES first for non-standard blog locations
                override_urls = BLOG_OVERRIDES.get(name, [])
                blog_found = False
                if override_urls:
                    log.debug(f"  {name}: checking BLOG_OVERRIDES: {override_urls}")
                for override_url in override_urls:
                    o_parsed = urlparse(override_url)
                    o_domain = o_parsed.netloc
                    o_path = o_parsed.path.rstrip("/")
                    log.debug(f"  {name}: trying override RSS on {o_domain}{o_path}")
                    # Try RSS on the override domain/path
                    blog_found = self._fetch_blog_rss(name, o_domain, o_parsed.scheme or "https", cutoff, out,
                                                       extra_blog_paths=[f"{o_path}/feed", f"{o_path}/rss", f"{o_path}/feed.xml"])
                    # If no RSS found, try HTML scrape on the override URL
                    if not blog_found:
                        log.debug(f"  {name}: no RSS at override, trying HTML scrape on {o_domain}{o_path}")
                        self._scrape_blog_html(name, o_domain, o_parsed.scheme or "https", cutoff, out,
                                                extra_paths=[o_path or "/"])
                        blog_found = len([s for s in out if s.meta.get("comp") == name and s.meta.get("type") == "blog"]) > 0
                    if blog_found:
                        blog_count = len([s for s in out if s.meta.get("comp") == name and s.meta.get("type") == "blog"])
                        log.info(f"  {name}: found {blog_count} blog posts via override ({override_url})")
                        break

                # Fall back to standard blog discovery on the main domain
                if not blog_found:
                    log.debug(f"  {name}: trying standard blog RSS on {domain}")
                    blog_found = self._fetch_blog_rss(name, domain, scheme, cutoff, out)
                if not blog_found:
                    log.debug(f"  {name}: trying standard HTML blog scrape on {domain}")
                    self._scrape_blog_html(name, domain, scheme, cutoff, out)

                # === PART C: Sitemap XML === (disabled in v5.0 — too slow, blog RSS is sufficient)
                # self._fetch_sitemaps(name, domain, scheme, cutoff, out)

                time.sleep(0.5)
            except Exception as e:
                log.warning(f"Sitemap {name}: {e}")
        log.info(f"Found {len(out)} blog posts total from competitor sites.")
        return out

    def _fetch_blog_rss(self, name, domain, scheme, cutoff, out, extra_blog_paths=None):
        """Try common blog RSS feed URLs. Returns True if any blog posts found."""
        found_any = False
        # Build RSS URL list: extra paths first, then standard paths
        all_paths = list(extra_blog_paths or []) + list(self.BLOG_PATHS)
        rss_urls = [f"{scheme}://{domain}{p}" for p in all_paths]
        # Try HTML <link rel="alternate"> discovery on /blog
        discovery_paths = ["/blog","/blog/","/en/blog","/news","/articles"]
        if extra_blog_paths:
            discovery_paths = [p.rstrip("/") for p in extra_blog_paths if p] + discovery_paths
        for blog_path in discovery_paths:
            try:
                r = GET(f"{scheme}://{domain}{blog_path}", timeout=8)
                if not r: continue
                soup = BeautifulSoup(r.text, "html.parser")
                for link in soup.find_all("link", rel="alternate"):
                    href = link.get("href","")
                    lt = link.get("type","")
                    if href and ("rss" in lt or "atom" in lt or "xml" in lt):
                        if href.startswith("/"): href = f"{scheme}://{domain}{href}"
                        if href not in rss_urls: rss_urls.insert(0, href)
                break  # Only need to check one blog page for <link>
            except Exception as e:
                log.debug(f"Blog RSS discovery {name} {blog_path}: {e}"); continue

        for rss_url in rss_urls:
            try:
                feed = feedparser.parse(rss_url)
                if not feed.entries: continue
                for e in feed.entries[:15]:
                    title = e.get("title","").strip()
                    link = e.get("link","")
                    pub = e.get("published", e.get("updated",""))
                    if not title: continue
                    if not recent(pub, 14): continue
                    pub_short = pub[:10] if pub else DATE
                    # Try to get a clean date
                    for fmt in ['%a, %d %b %Y %H:%M:%S %z','%Y-%m-%dT%H:%M:%S%z','%Y-%m-%d']:
                        try:
                            pub_short = datetime.strptime(pub.replace('GMT','+0000'),fmt).strftime("%Y-%m-%d")
                            break
                        except (ValueError, TypeError): continue
                    out.append(Signal("sitemap", f"{name}: {title[:80]}",
                        f"Blog post on {name} ({pub_short})",
                        url=link, score=50,
                        meta={"comp":name,"lastmod":pub_short,"type":"blog","cats":cats(title),
                              "biz_cat":biz_cats(title)[0],"activity_type":"blog_post"}))
                    found_any = True
                if found_any:
                    log.info(f"Blog RSS {name}: found posts")
                    break  # Got a working feed, no need to try more
            except Exception as e:
                log.debug(f"Blog RSS {name} feed: {e}"); continue
        return found_any

    def _scrape_blog_html(self, name, domain, scheme, cutoff, out, extra_paths=None):
        """Fallback: scrape the blog HTML page for article links & titles."""
        paths = list(extra_paths or []) + ["/blog","/blog/","/en/blog","/news","/articles","/hub","/hub/"]
        for blog_path in paths:
            try:
                r = GET(f"{scheme}://{domain}{blog_path}", timeout=10)
                if not r or r.status_code != 200: continue
                soup = BeautifulSoup(r.text, "html.parser")
                # Look for article-like elements
                articles = soup.find_all(["article","div"],
                    class_=re.compile(r"post|article|blog|card|entry", re.I))
                if not articles:
                    # Fallback: just find all links that look like blog posts
                    articles = soup.find_all("a", href=re.compile(r"/blog/|/post/|/article/|/news/|/hub/"))
                found = 0
                for art in articles[:25]:
                    # Get title
                    h = art.find(["h1","h2","h3","h4"])
                    if h:
                        title = h.get_text(strip=True)
                    elif art.name == "a":
                        title = art.get_text(strip=True)
                    else:
                        continue
                    if not title or len(title) < 10: continue
                    # Get link
                    link_el = art.find("a", href=True) if art.name != "a" else art
                    href = link_el.get("href","") if link_el else ""
                    if href.startswith("/"): href = f"{scheme}://{domain}{href}"
                    # Get date if available
                    time_el = art.find("time")
                    date_str = time_el.get("datetime","")[:10] if time_el else ""
                    if date_str:
                        try:
                            if datetime.strptime(date_str, "%Y-%m-%d") < cutoff: continue
                        except ValueError: pass
                    out.append(Signal("sitemap", f"{name}: {title[:80]}",
                        f"Blog page on {name}" + (f" ({date_str})" if date_str else ""),
                        url=href, score=45,
                        meta={"comp":name,"lastmod":date_str or "recent","type":"blog","cats":cats(title),
                              "biz_cat":biz_cats(title)[0],"activity_type":"blog_post"}))
                    found += 1
                if found:
                    log.info(f"Blog scrape {name}: {found} posts")
                    break
            except Exception as e:
                log.debug(f"Blog scrape {name} {blog_path}: {e}"); continue

    def _fetch_sitemaps(self, name, domain, scheme, cutoff, out):
        """Fetch XML sitemaps for new/modified pages."""
        sitemap_urls = []
        # Step 1: robots.txt
        try:
            r_robots = GET(f"{scheme}://{domain}/robots.txt", timeout=8)
            if r_robots:
                for line in r_robots.text.splitlines():
                    if line.lower().startswith("sitemap:"):
                        sm_url = line.split(":",1)[1].strip()
                        if sm_url and sm_url not in sitemap_urls:
                            sitemap_urls.append(sm_url)
        except Exception as e: log.debug(f"Sitemap robots.txt {domain}: {e}")
        # Step 2: Common patterns
        sitemap_urls.extend([
            f"{scheme}://{domain}/sitemap.xml",
            f"{scheme}://{domain}/sitemap_index.xml",
            f"{scheme}://{domain}/sitemap-pages.xml",
            f"{scheme}://{domain}/page-sitemap.xml",
            f"{scheme}://{domain}/post-sitemap.xml",
            f"{scheme}://{domain}/wp-sitemap.xml",
            f"{scheme}://{domain}/sitemap1.xml",
        ])
        sitemap_urls = list(dict.fromkeys(sitemap_urls))
        found_urls = []
        for sm_url in sitemap_urls:
            try:
                r = GET(sm_url, timeout=10)
                if not r: continue
                root = ET.fromstring(r.content)
                ns = {"s":"http://www.sitemaps.org/schemas/sitemap/0.9"}
                sitemaps = root.findall("s:sitemap", ns)
                if sitemaps:
                    for sm in sitemaps[:5]:
                        loc = sm.find("s:loc", ns)
                        if loc is not None and loc.text:
                            try:
                                r2 = GET(loc.text.strip(), timeout=10)
                                if r2: found_urls.extend(self._parse_urlset(r2.content, ns, cutoff))
                            except Exception as e:
                                log.debug(f"Sitemap sub {name}: {e}"); continue
                else:
                    found_urls.extend(self._parse_urlset(r.content, ns, cutoff))
                if found_urls: break
            except Exception as e:
                log.debug(f"Sitemap XML {name}: {e}"); continue
        log.info(f"Sitemap {name}: {len(found_urls)} new pages")
        # Deduplicate against existing signals for this competitor
        existing_urls = set()
        for s in out:
            if s.meta.get("comp") == name: existing_urls.add(s.url)
        # For competitors with blog overrides, skip generic category/landing pages from sitemaps
        # (we already have their blog content via the override)
        has_override = name in BLOG_OVERRIDES
        for page_url, lastmod in found_urls[:20]:
            if page_url in existing_urls: continue
            path = urlparse(page_url).path.strip("/")
            page_title = path.split("/")[-1].replace("-"," ").replace("_"," ").title() if path else page_url
            # Classify page type by URL path
            act_type = "new_page"
            if any(x in page_url.lower() for x in ["/blog/","/post/","/article/","/news/","/hub/news/"]): act_type = "blog_post"
            elif any(x in page_url.lower() for x in ["/promo","/deal","/sale","/offer"]): act_type = "promotion"
            elif any(x in page_url.lower() for x in ["/product","/item","/gift-card","/game"]): act_type = "product_page"
            # Skip generic sitemap pages for competitors where we already have blog content
            if has_override and act_type == "new_page": continue
            out.append(Signal("sitemap", f"{name}: {page_title[:80]}",
                f"New/updated page on {name} ({lastmod})",
                url=page_url, score=40,
                meta={"comp":name,"lastmod":lastmod,"type":"page","cats":cats(page_title),
                      "biz_cat":biz_cats(page_title)[0],"activity_type":act_type}))

    def _parse_urlset(self, content, ns, cutoff):
        results = []
        try:
            root = ET.fromstring(content)
            for url_el in root.findall("s:url", ns):
                loc = url_el.find("s:loc", ns)
                mod = url_el.find("s:lastmod", ns)
                if loc is None: continue
                loc_text = loc.text.strip() if loc.text else ""
                if mod is None: continue
                mod_text = mod.text.strip()[:10] if mod.text else ""
                try:
                    mod_date = datetime.strptime(mod_text, "%Y-%m-%d")
                    if mod_date >= cutoff:
                        results.append((loc_text, mod_text))
                except ValueError: continue
        except Exception as e: log.debug(f"Sitemap parse: {e}")
        return results

# =============================================================================
# SECTION 4B - HISTORY & TRENDS (week-over-week comparison)
# =============================================================================

def load_previous_history():
    """Load the most recent history JSON file for week-over-week comparison."""
    import glob as _glob
    history_files = sorted(_glob.glob("history_*.json"), reverse=True)
    for hf in history_files:
        try:
            date_str = hf.replace("history_", "").replace(".json", "")
            if date_str == DATE: continue
            with open(hf) as f: return json.load(f)
        except Exception as e: log.debug(f"Failed to load {hf}: {e}")
    return None

def compute_trends(current_cands, all_sig, previous_history):
    """Compute week-over-week changes for KPIs and candidate scores."""
    if not previous_history:
        return {"kpi_deltas": {}, "movers": [], "new_entries": [], "dropped": []}
    prev_kpis = previous_history.get("kpis", {})
    curr_kpis = {
        "total_signals": sum(len(v) for v in all_sig.values()),
        "candidates": len(current_cands),
        "multi_source": len([c for c in current_cands if c.sources >= 2]),
    }
    kpi_deltas = {k: curr_kpis.get(k, 0) - prev_kpis.get(k, 0) for k in curr_kpis}
    prev_by_title = {c["title"]: c for c in previous_history.get("top_candidates", [])}
    movers, new_entries = [], []
    for c in current_cands[:50]:
        prev = prev_by_title.get(c.title)
        if prev:
            delta = c.score - prev["score"]
            if abs(delta) > 3:
                movers.append({"title": c.title, "score": c.score, "prev_score": prev["score"], "delta": round(delta, 1)})
        else:
            new_entries.append({"title": c.title, "score": c.score})
    curr_titles = {c.title for c in current_cands[:50]}
    dropped = [{"title": c["title"], "prev_score": c["score"]}
               for c in previous_history.get("top_candidates", [])[:30]
               if c["title"] not in curr_titles]
    return {
        "kpi_deltas": kpi_deltas,
        "movers": sorted(movers, key=lambda x: -abs(x["delta"]))[:10],
        "new_entries": new_entries[:10],
        "dropped": dropped[:10],
    }

# =============================================================================
# SECTION 5 - CONCURRENT ORCHESTRATOR
# =============================================================================

FETCH_TIMEOUT = 90  # seconds per fetcher

def fetch_all():
    fetchers = {
        "trends":TrendsFetcher(),"reddit":RedditFetcher(),"steam":SteamFetcher(),
        "wiki":WikiFetcher(),"youtube":YTFetcher(),"news":NewsFetcher(),
        "oxylabs_news":OxylabsNewsFetcher(),
        "competitor":CompetitorFetcher(),"cheapshark":CheapSharkFetcher(),
        "steamspy":SteamSpyFetcher(),"gamerpower":GamerPowerFetcher(),
        "epic":EpicFreeFetcher(),
        "gog":GOGFetcher(),"humble":HumbleFetcher(),
        "freetogame":FreeToGameFetcher(),"anime":AnimeFetcher(),
        "sitemap":SitemapFetcher(),
    }
    log.info(f"  Launching {len(fetchers)} data collectors in parallel (6 at a time)...")
    for name in sorted(fetchers.keys()):
        log.debug(f"  Queued: {name}")
    print("\n" + "="*60); print(f"FETCHING {len(fetchers)} SOURCES (concurrent)"); print("="*60)
    results = {}
    with ThreadPoolExecutor(max_workers=6) as ex:
        futs = {ex.submit(f.fetch):n for n,f in fetchers.items()}
        try:
            for fut in as_completed(futs, timeout=420):
                n = futs[fut]
                try:
                    results[n] = fut.result(timeout=FETCH_TIMEOUT)
                    log.debug(f"  Done: {n} -> {len(results[n])} items")
                except TimeoutError: log.warning(f"  {n} took too long, skipping"); results[n] = []
                except Exception as e: log.warning(f"  {n} ran into an error: {e}"); results[n] = []
        except TimeoutError:
            for fut,n in futs.items():
                if n not in results:
                    log.warning(f"  {n} was still running when time ran out, skipping")
                    results[n] = []
    total = sum(len(v) for v in results.values())
    for n,s in sorted(results.items()): print(f"  {n}: {len(s)}")
    failed = [n for n, s in results.items() if len(s) == 0]
    if failed: log.info(f"  Sources that returned nothing: {', '.join(failed)}")
    print(f"  TOTAL: {total}"); return results

# =============================================================================
# SECTION 6 - DEDUP & MERGE
# =============================================================================

_ROMANS = {'vi':'6','vii':'7','viii':'8','ix':'9','iv':'4','iii':'3','ii':'2','xl':'40','xx':'20','xv':'15','x':'10'}

def _normalize_title(text):
    """Aggressive title normalization for dedup: lowercase, strip punctuation, convert roman numerals."""
    lo = re.sub(r'[^a-z0-9 ]','',text.lower())
    words = lo.split()
    return ' '.join(_ROMANS.get(w, w) for w in words)

def _merge_into_cands(sigs, cands):
    """Merge signals into candidates using best-match fuzzy dedup with token pre-filter."""
    for sig in sigs:
        best_cand = None; best_ratio = 0.0
        a_norm = _normalize_title(sig.title)
        if not a_norm:
            cands.append(Candidate(title=sig.title, signals=[sig])); continue
        a_tokens = set(a_norm.split())
        nums_a = re.findall(r'\d+', a_norm)
        for c in cands:
            b_norm = _normalize_title(c.title)
            if not b_norm: continue
            nums_b = re.findall(r'\d+', b_norm)
            if nums_a and nums_b and nums_a != nums_b: continue
            # Cheap pre-filter: token overlap (Jaccard)
            b_tokens = set(b_norm.split())
            union = a_tokens | b_tokens
            jaccard = len(a_tokens & b_tokens) / len(union) if union else 0
            if jaccard < 0.25: continue
            r = SequenceMatcher(None, a_norm, b_norm).ratio()
            if r >= FUZZ_T and r > best_ratio:
                best_ratio = r; best_cand = c
        if best_cand:
            best_cand.signals.append(sig)
        else:
            cands.append(Candidate(title=sig.title, signals=[sig]))

def dedup(all_sig):
    flat = [s for sigs in all_sig.values() for s in sigs]
    if not flat: return []
    # Group by primary category for blocking (reduces O(n^2))
    by_cat = defaultdict(list); no_cat = []
    for sig in flat:
        sig_cats = sig.meta.get("cats", [])
        if sig_cats: by_cat[sig_cats[0]].append(sig)
        else: no_cat.append(sig)
    cands = []
    for cat, sigs in by_cat.items():
        _merge_into_cands(sigs, cands)
    _merge_into_cands(no_cat, cands)
    # Assign categories, sources, biz_category
    for c in cands:
        src_types = set(); all_cats = []
        for s in c.signals:
            st = s.source
            if st in ("news","google"): st = "news"
            src_types.add(st)
            all_cats.extend(s.meta.get("cats",[]))
        c.sources = len(src_types)
        cat_counts = Counter(all_cats)
        c.categories = [cat for cat,_ in cat_counts.most_common()]
        c.category = c.categories[0] if c.categories else "General"
        biz_counts = Counter(BIZ_CATS.get(cat,"GMG") for cat in c.categories)
        c.biz_categories = [bc for bc,_ in biz_counts.most_common()]
        c.biz_category = c.biz_categories[0] if c.biz_categories else "GMG"
        c.url = _best_url(c.signals)
    print(f"  {len(flat)} signals -> {len(cands)} candidates"); return cands

# =============================================================================
# SECTION 7 - COMPOSITE SCORING  (BUG FIX: removed erroneous *100)
# =============================================================================

def comp_score(cands):
    for c in cands:
        by_src = defaultdict(list)
        for s in c.signals:
            st = s.source
            if st in ("news","google"): st = "news"
            by_src[st].append(s.score)
        ws = 0.0; tw = 0.0
        for st,scores in by_src.items():
            w = W.get(st,0.04)
            ws += (sum(scores)/len(scores))*w; tw += w
        base = ws/tw if tw > 0 else 0
        mult = CONF.get(c.sources, CONF_DEFAULT)
        c.score = round(min(base * mult, 100), 1)
    cands.sort(key=lambda x: -x.score)
    return cands

def get_top3_per_biz_cat(cands, opps):
    """For each business category (GMG/ENT/PPM/MTU), return top 3 items.
    ALWAYS uses keyword-matched biz_category from candidates, NOT the AI's assignment
    (the AI often misclassifies e.g. Fortnite as MTU instead of GMG)."""
    top3 = {"GMG": [], "ENT": [], "PPM": [], "MTU": []}
    # First pull from AI opportunities (already prioritized)
    for o in opps:
        title = o.get("title", "")
        mc = _match_cand(title, cands)
        # ALWAYS prefer the keyword-matched category from the candidate
        if mc:
            bc = mc.biz_category
        else:
            bc = biz_cats(title)[0]
        if bc not in top3: bc = "GMG"
        if len(top3[bc]) < 3:
            top3[bc].append({
                "title": title,
                "why_now": o.get("why_now", ""),
                "urgency": o.get("urgency", "medium"),
                "revenue_signal": o.get("revenue_signal", ""),
                "score": mc.score if mc else 0,
                "url": mc.url if mc else "",
                "category": mc.category if mc else o.get("category", ""),
            })
    # Fill remaining slots from scored candidates
    for bc in top3:
        if len(top3[bc]) < 3:
            used = {item["title"] for item in top3[bc]}
            for c in cands:
                if c.biz_category == bc and c.title not in used:
                    top3[bc].append({
                        "title": c.title,
                        "why_now": f"Across {c.sources} sources",
                        "urgency": "high" if c.score > 50 else "medium",
                        "revenue_signal": "",
                        "score": c.score,
                        "url": c.url,
                        "category": c.category,
                    })
                    if len(top3[bc]) >= 3: break
    return top3

# =============================================================================
# SECTION 8 - 4-PASS AI
# =============================================================================

def _gemini_json(prompt, retries=2):
    for i in range(retries+1):
        try:
            GEMINI_LIMITER.wait()
            r = GCLIENT.models.generate_content(model=GEMINI_MODEL,contents=prompt,
                config=gtypes.GenerateContentConfig(response_mime_type="application/json",temperature=0.2))
            return json.loads(r.text.strip())
        except json.JSONDecodeError:
            if i < retries: time.sleep(2)
        except Exception as e:
            log.warning(f"    AI request attempt {i+1} failed: {e}")
            if i < retries: time.sleep(3)
    return None

def _gemini_grounded(prompt, retries=2):
    for i in range(retries+1):
        try:
            GEMINI_LIMITER.wait()
            r = GCLIENT.models.generate_content(model=GEMINI_MODEL,contents=prompt,
                config=gtypes.GenerateContentConfig(tools=[gtypes.Tool(google_search=gtypes.GoogleSearch())],temperature=0.3))
            sources = []
            try:
                gm = r.candidates[0].grounding_metadata
                if gm and gm.grounding_chunks:
                    for ch in gm.grounding_chunks:
                        if ch.web: sources.append({"url":ch.web.uri,"title":ch.web.title})
            except Exception as e: log.debug(f"Grounding metadata: {e}")
            return {"text":r.text,"sources":sources}
        except Exception as e:
            log.warning(f"    AI fact-check attempt {i+1} failed: {e}")
            if i < retries: time.sleep(3)
    return None

def pass0_ground(cands, events):
    print("\nAI PASS 0: Real-time intelligence (Google Search grounding)...")
    topics = [c.title for c in cands[:15]]
    urgent = [e["name"] for e in events[:10] if e.get("urgency") in ("critical","high")]
    prompt = f"""You are a market intelligence analyst for Recharge.com, which sells gaming gift cards, streaming subscriptions, and digital game credits online.

TODAY: {NOW.strftime('%B %d, %Y')}

TASK: Search the web for real-time developments on these trending topics and upcoming events. For each topic with significant recent news, provide a structured assessment.

TRENDING TOPICS (from our data pipeline):
{chr(10).join(f'- {t}' for t in topics)}

IMMINENT EVENTS:
{chr(10).join(f'- {e}' for e in urgent)}

For each topic with notable developments in the last 48 hours, report:
1. WHAT HAPPENED: Specific facts with dates (do not speculate)
2. REVENUE ANGLE: How does this create a gift card/top-up buying opportunity?
3. SEARCH DEMAND: What are consumers searching for right now related to this?
4. TIMING: Is this a 24-hour window, this-week opportunity, or multi-week trend?

RULES:
- Only report topics with verifiable recent developments
- Skip anything without concrete news from the last 48 hours
- Be specific: include numbers, dates, platform names, and pricing where available
- Prioritize events that directly drive digital gift card or top-up purchases"""
    result = _gemini_grounded(prompt)
    if result: print(f"  Got real-time intel ({len(result.get('sources',[]))} sources)"); return result
    print("  Grounding unavailable"); return {"text":"","sources":[]}

def pass1(cands, events, ground_intel, all_sig=None):
    print("AI PASS 1: Prioritize...")
    all_sig = all_sig or {}

    # Split signals into TRENDING (time-sensitive) vs MARKET (static background data)
    TRENDING_SOURCES = {"news","oxylabs_news","reddit","youtube","trends"}
    trending_titles = []
    for src in TRENDING_SOURCES:
        for s in all_sig.get(src,[]):
            trending_titles.append(f"- [{s.source.upper()}] {s.title} | {s.desc}")
    # Deduplicate trending titles
    seen_t = set()
    unique_trending = []
    for t in trending_titles:
        k = re.sub(r'[^a-z0-9]','',t[:60].lower())
        if k not in seen_t: seen_t.add(k); unique_trending.append(t)

    trending_text = "\n".join(unique_trending[:80]) if unique_trending else "No trending signals this run."
    market_text = "\n".join(f"{i+1}. [{c.score}] {c.title} (sources={c.sources}, cat={c.category})" for i,c in enumerate(cands[:20]))
    et = "\n".join(f"- {e['name']} ({e['category']}): {e['status']} - {e['description']}" for e in events[:15])
    intel = ground_intel.get("text","")[:2000]

    # Build numbered reference list of all available titles
    ref_titles = []; seen_ref = set()
    for t in unique_trending:
        title = t.split("] ", 1)[-1].split(" | ")[0].strip()
        if title not in seen_ref: seen_ref.add(title); ref_titles.append(title)
    for c in cands[:20]:
        if c.title not in seen_ref: seen_ref.add(c.title); ref_titles.append(c.title)
    ref_text = "\n".join(f"REF-{i+1}: {t}" for i, t in enumerate(ref_titles[:50]))

    prompt = f"""You are a senior growth strategist at Recharge.com (gaming gift cards, streaming subscriptions, digital credits).
TODAY: {NOW.strftime('%B %d, %Y')}

=== SECTION A: THIS WEEK'S TRENDING NEWS (from RSS, Reddit, YouTube, Google Trends) ===
These are REAL headlines and posts from the past 7 days:
{trending_text}

=== SECTION B: REAL-TIME INTELLIGENCE (Google Search, verified today) ===
{intel if intel else "Not available"}

=== SECTION C: UPCOMING EVENTS ===
{et}

=== SECTION D: MARKET DATA (background, always-popular games/services) ===
{market_text}

=== REFERENCE LIST (you MUST use titles exactly as they appear) ===
{ref_text}

YOUR TASK: Pick TOP 15 opportunities for Recharge.com to act on THIS WEEK.

CRITICAL RULES:
1. AT LEAST 10 of your 15 picks MUST come from Section A (trending news) or Section C (events). These are things actually happening NOW.
2. At most 5 picks can come from Section D (market data), and ONLY if they have a specific time-sensitive angle.
3. NEVER pick a generic topic like "Roblox" or "Fortnite" unless there is a specific news headline about it in Section A.
4. Every pick MUST answer: "What happened THIS WEEK?" — if you can't answer that, don't include it.
5. Use EXACT titles from the REFERENCE LIST above. Reference by REF number. Do NOT rephrase or invent titles.
6. Spread across categories. Include at least 2 picks per business category (GMG/ENT/PPM/MTU) if possible.

Return JSON: {{"opportunities": [
  {{"title": "exact title from reference list", "category": "...", "biz_category": "GMG|ENT|PPM|MTU",
    "urgency": "critical|high|medium",
    "confidence": 0.0-1.0, "why_now": "what happened THIS WEEK", "revenue_signal": "how this drives gift card purchases"}}
]}}"""
    r = _gemini_json(prompt)
    if r and "opportunities" in r: print(f"  {len(r['opportunities'])} opportunities"); return r["opportunities"]
    return [{"title":c.title,"category":c.category,"urgency":"high" if c.score>50 else "medium",
             "confidence":min(c.score/100,1.0),"why_now":f"Across {c.sources} sources",
             "revenue_signal":"Multiple signals indicate purchase intent"} for c in cands[:15]]

def pass3(opps, ground_intel):
    print("AI PASS 2: Executive synthesis (grounded)...")
    prompt = f"""You are the VP of Growth presenting to C-suite at Recharge.com.

CONTEXT: Recharge.com sells digital gift cards, game credits, streaming subscriptions, and mobile top-up credits to consumers worldwide.

THIS WEEK'S VERIFIED OPPORTUNITIES (AI-prioritized):
{json.dumps(opps[:10], indent=2)}

REAL-TIME INTELLIGENCE:
{ground_intel.get("text","")[:2000] if isinstance(ground_intel, dict) else ""}

Write a razor-sharp executive briefing. No filler. Every word must be actionable.

Return JSON:
{{"summary": "2-3 sentences. Lead with the #1 revenue opportunity and its estimated impact. Include specific product names (e.g., 'PlayStation Store cards' not 'gift cards').",
"actions": [
  "Commercial Team: [specific action] - [specific deadline, e.g., 'by Wednesday']",
  "Content Team: [specific action] - [specific deadline]",
  "Marketing Team: [specific action] - [specific deadline]",
  "Product Team: [specific action if applicable] - [specific deadline]"
],
"predictions": ["2-3 specific, testable predictions for the next 7-14 days. Include what product categories will spike and why."],
"risks": ["2-3 specific risks with mitigation suggestions. Include competitor threats."],
"category_highlights": {{
  "GMG": "One sentence: the single most important gaming takeaway with a specific product recommendation",
  "ENT": "One sentence: entertainment takeaway with specific streaming service",
  "PPM": "One sentence: prepaid takeaway",
  "MTU": "One sentence: mobile top-up takeaway"
}}}}

CRITICAL: Each 'actions' entry must be a plain STRING. Do NOT nest objects."""
    r = _gemini_json(prompt)
    if r and "summary" in r: print("  Done"); return r
    return {"summary":"Multiple revenue opportunities identified.","actions":["Commercial team: act on top opportunities",
            "Content team: update landing pages","Marketing team: monitor competitors"],
            "predictions":["Watch for major updates"],"risks":["Competitor pricing"]}

def _fetch_competitor_news():
    """Use Oxylabs to get real-time news about each competitor."""
    if not OXYLABS_USER or not OXYLABS_PASS: return {}
    comp_news = {}
    targets = list(COMPETITORS.keys()) + [c for c in SITEMAP_COMPETITORS if c not in COMPETITORS]
    for name in targets:
        try:
            payload = {
                "source": "google_search", "query": f'"{name}" gift card OR gaming OR digital OR top-up',
                "parse": True, "context": [{"key":"tbm","value":"nws"},{"key":"tbs","value":"qdr:w"}],
                "geo_location": "United States", "locale": "en-US", "limit": 5,
            }
            r = requests.post("https://realtime.oxylabs.io/v1/queries",
                auth=(OXYLABS_USER, OXYLABS_PASS), json=payload, timeout=30)
            if r.status_code != 200: continue
            data = r.json()
            results = data.get("results", [])
            if not results: continue
            content = results[0].get("content", {})
            if not isinstance(content, dict): continue
            items = content.get("results", {}).get("main", [])
            comp_news[name] = [{"title":it.get("title",""),"source":it.get("source",""),
                "age":it.get("relative_publish_date",""),"url":it.get("url",""),
                "desc":it.get("desc","")} for it in items[:5] if it.get("title")]
            time.sleep(0.3)
        except Exception as e: log.debug(f"Competitor news {name}: {e}"); continue
    return comp_news

def pass_competitor(comp_signals, sitemap_signals):
    """Advanced AI competitor intelligence with real-time news, sitemap analysis, and strategic assessment."""
    print("AI PASS: Competitor deep intelligence...")
    # 1. Gather all raw competitor data
    comp_data = []
    for s in (comp_signals or []) + (sitemap_signals or []):
        comp_data.append(f"- {s.title}: {s.desc} ({s.url})")
    # 2. Fetch real-time competitor news via Oxylabs
    print("  Fetching real-time competitor news...")
    comp_news = _fetch_competitor_news()
    news_text = ""
    for name, articles in comp_news.items():
        if articles:
            news_text += f"\n{name} IN THE NEWS:\n"
            for a in articles:
                news_text += f"  - [{a['age']}] {a['title']} (via {a['source']})\n"
    total_news = sum(len(v) for v in comp_news.values())
    print(f"  Got {total_news} competitor news articles")
    # 3. Build sitemap stats per competitor
    sitemap_stats = defaultdict(lambda: {"blogs":0,"pages":0,"promos":0,"titles":[]})
    for s in (sitemap_signals or []):
        cn = s.meta.get("comp","?"); at = s.meta.get("activity_type","new_page")
        if at == "blog_post": sitemap_stats[cn]["blogs"] += 1
        elif at == "promotion": sitemap_stats[cn]["promos"] += 1
        else: sitemap_stats[cn]["pages"] += 1
        sitemap_stats[cn]["titles"].append(s.title.replace(f"{cn}: ","",1)[:60])
    stats_text = ""
    for cn, st in sitemap_stats.items():
        stats_text += f"\n{cn} BLOG ACTIVITY: {st['blogs']} blog posts"
        if st["titles"]: stats_text += f"\n  Recent: {', '.join(st['titles'][:5])}"

    if not comp_data and not news_text:
        print("  No competitor data"); return []

    prompt = f"""You are a senior competitive intelligence analyst at Recharge.com (digital gift cards, game credits, streaming subs, mobile top-up).

=== COMPETITOR HOMEPAGE & DEALS SCRAPING ===
{chr(10).join(comp_data[:40]) if comp_data else "No scraping data available"}

=== COMPETITOR NEWS COVERAGE (real-time from Google News) ===
{news_text if news_text else "No news coverage found"}

=== COMPETITOR BLOG ACTIVITY ===
{stats_text if stats_text else "No sitemap changes detected"}

OUR DIRECT COMPETITORS: G2A, Eneba, CDKeys, Kinguin (game key marketplaces)
OUR INDIRECT COMPETITORS: Dundle, MobileRecharge, TalkHome, OnTopUp, KarteDirekt, Aufladen (gift cards, mobile top-up)

Create a COMPREHENSIVE competitor intelligence report. For each competitor with data:

Return JSON: {{"competitors": [
  {{
    "name": "competitor name",
    "type": "direct|indirect",
    "headline": "one punchy sentence summarizing what they did this week",
    "strategy_signal": "what their activity TELLS US about their strategy (1 sentence)",
    "activities": ["3-5 specific activities with context"],
    "news_highlights": ["top 1-2 news stories about them, if any"],
    "website_moves": "summary of their blog/page changes",
    "threat_level": "high|medium|low",
    "threat_reason": "WHY this threat level (1 sentence)",
    "our_response": "what Recharge.com should do about this competitor (1 specific action)",
    "notable": "the single most interesting/concerning thing they did"
  }}
],
"market_summary": "2-3 sentences on overall competitive landscape this week",
"biggest_threat": "which competitor is the biggest threat right now and why (1 sentence)",
"opportunity_gap": "what are competitors NOT doing that we could exploit (1 sentence)"
}}
Be specific. Use real data. No generic statements."""
    r = _gemini_json(prompt)
    if r and "competitors" in r:
        print(f"  {len(r['competitors'])} competitors analyzed (deep)")
        return r
    # Fallback
    by_comp = defaultdict(list)
    for s in (comp_signals or []) + (sitemap_signals or []):
        cn = s.meta.get("comp","Unknown"); by_comp[cn].append(s.title)
    return {"competitors":[{"name":cn,"type":"direct" if cn in COMPETITORS else "indirect",
             "headline":f"{len(acts)} activities detected","strategy_signal":"Active this week",
             "activities":acts[:5],"news_highlights":[],"website_moves":"See sitemap data",
             "threat_level":"medium","threat_reason":"Active competitor","our_response":"Monitor",
             "notable":acts[0] if acts else ""} for cn,acts in by_comp.items()],
            "market_summary":"Competitors are active. Check details.","biggest_threat":"","opportunity_gap":""}

def pass_newsletter(opps, top3, comp_intel, events, executive):
    """Generate witty, engaging copy for the email newsletter."""
    print("AI PASS: Newsletter copy...")
    prompt = f"""You write an internal company newsletter for Recharge.com (digital gift cards, game credits, streaming subs, mobile top-up).
Your tone: knowledgeable but fun. Like a smart colleague who's really into gaming and entertainment.
Think: The Morning Brew meets gaming news. Entertaining but informative.

TOP STORIES: {json.dumps(opps[:5], default=str)}
CATEGORY TOP 3: {json.dumps(top3, default=str)}
COMPETITOR INTEL: {json.dumps(comp_intel, default=str) if comp_intel else "None"}
EXECUTIVE SUMMARY: {executive.get("summary", "") if isinstance(executive, dict) else ""}

Generate newsletter copy. Return JSON with these fields:
1. "big_picture": 2 punchy sentences about this week (NOT corporate-speak, be specific)
2. "hot_items": array of 3 objects, each with "title" (catchy, max 10 words) and "blurb" (1 witty sentence, max 20 words)
3. "category_blurbs": object where key=GMG/ENT/PPM/MTU, value=array of 3 objects with "title" and "blurb" (max 15 words each)
4. "competitor_watch": 2-3 sentences about what competitors did, written like casual intel/gossip
5. "dont_miss": array of max 3 objects with "event" and "why_care" (fun one-liner, max 15 words)
6. "crystal_ball": 1-2 fun prediction sentences about next week

Keep ALL blurbs VERY short. Be specific with game/service names, not generic. Have personality."""
    r = _gemini_json(prompt)
    if r and "big_picture" in r: print("  Newsletter copy generated"); return r
    # Fallback
    print("  Using fallback copy")
    return {
        "big_picture": executive.get("summary","Another week of opportunities in gaming and entertainment.") if isinstance(executive,dict) else "Another big week.",
        "hot_items": [{"title":o.get("title","")[:40],"blurb":o.get("why_now","") or "Trending this week"} for o in opps[:3]],
        "category_blurbs": {bc: [{"title":item["title"][:40],"blurb":item.get("why_now","") or "Trending"} for item in items[:3]] for bc,items in top3.items()},
        "competitor_watch": "Competitors are active this week. Check the dashboard for details.",
        "dont_miss": [{"event":e["name"],"why_care":e.get("description","")} for e in events[:3] if e.get("urgency") in ("critical","high")],
        "crystal_ball": "Keep an eye on upcoming releases and seasonal events.",
    }

def _validate_ai_output(data, required_keys, list_key=None):
    """Validate AI JSON output has expected structure."""
    if not isinstance(data, dict): return False
    for k in required_keys:
        if k not in data: return False
    if list_key and not isinstance(data.get(list_key), list): return False
    return True

def run_ai(cands, events, comp_sigs, all_sig=None):
    print("\n" + "="*60); print("AI ANALYSIS (with Google Search grounding)"); print("="*60)
    log.info("  I'm reviewing all the data now. Here's my 4-pass approach:")
    log.info(f"  I have {len(cands)} opportunities, {len(events)} events, and {len(comp_sigs)} competitor signals to work with.")

    # Pass 0: Grounding (non-critical)
    log.info("")
    log.info("  Pass 1 of 4 — I'll verify facts by searching Google for the latest news...")
    t_p0 = time.time()
    try:
        ground = pass0_ground(cands, events)
        log.info(f"    Done in {time.time()-t_p0:.1f}s — I found {len(ground.get('sources',[]))} verified sources.")
    except Exception as e:
        log.warning(f"    Couldn't complete fact-checking ({e}), but that's okay.")
        log.warning(f"    I'll continue without it — the other passes will still work fine.")
        ground = {"text":"","sources":[]}

    # Pass 1: Prioritize (critical)
    log.info("")
    log.info("  Pass 2 of 4 — I'll pick the most important priorities from all this data...")
    t_p1 = time.time()
    try:
        opps = pass1(cands, events, ground, all_sig or {})
        log.info(f"    Done in {time.time()-t_p1:.1f}s — I selected {len(opps)} top opportunities.")
        for o in opps[:5]:
            log.debug(f"    - {o.get('title','')[:50]} ({o.get('urgency','')})")
    except Exception as e:
        log.error(f"    Prioritization hit a snag ({e}) — I'll use score-based ranking as a fallback.")
        opps = [{"title":c.title,"category":c.category,"biz_category":c.biz_category,
                 "urgency":"high" if c.score>50 else "medium",
                 "confidence":min(c.score/100,1.0),"why_now":f"Across {c.sources} sources",
                 "revenue_signal":"Multiple signals"} for c in cands[:15]]

    # Pass 2 (exec) and Pass 3 (competitor) run in parallel
    log.info("")
    log.info("  Pass 3 of 4 — I'll write the executive brief for leadership...")
    log.info("  Pass 4 of 4 — I'll analyze competitors and their moves...")
    log.info("  (I'll run both at the same time to save time.)")
    t_p23 = time.time()
    ex = {"summary":"Analysis in progress.","actions":[],"predictions":[],"risks":[]}
    comp_intel = {"competitors":[],"market_summary":"","biggest_threat":"","opportunity_gap":""}
    with ThreadPoolExecutor(max_workers=2) as pool:
        fut_ex = pool.submit(pass3, opps, ground)
        fut_comp = pool.submit(pass_competitor, comp_sigs, (all_sig or {}).get("sitemap", []))
        try:
            ex = fut_ex.result(timeout=120)
            log.info(f"    Executive summary done — I recommended {len(ex.get('actions',[]))} actions.")
        except Exception as e: log.warning(f"    Executive summary couldn't be completed: {e}")
        try:
            comp_intel = fut_comp.result(timeout=120)
            log.info(f"    Competitor analysis done — I reviewed {len(comp_intel.get('competitors',[]))} competitors.")
        except Exception as e: log.warning(f"    Competitor analysis couldn't be completed: {e}")
    log.info(f"    Both passes finished in {time.time()-t_p23:.1f}s.")

    return {"opportunities":opps,"executive":ex,"ground_intel":ground,"competitor_intel":comp_intel}

# =============================================================================
# SECTION 9 - GOOGLE SHEETS
# =============================================================================

def write_sheets(cands, ai, events):
    if not IS_COLAB: print("\nGOOGLE SHEETS... skipped (not Colab)"); return None
    print("\nGOOGLE SHEETS...", end=" ")
    try:
        from google.colab import auth; auth.authenticate_user()
        creds,_ = _gauth(); gc = gspread.authorize(creds)
    except:
        if not HAS_GSPREAD: print("skipped"); return None
        try: gc = gspread.service_account()
        except: print("auth failed"); return None
    try: sh = gc.create(f"Recharge Scanner {DATE}")
    except Exception as e: print(f"failed: {e}"); return None
    try:
        ex = ai.get("executive",{})
        ws1 = sh.sheet1; ws1.update_title("Dashboard")
        ws1.update(values=[["RECHARGE OPPORTUNITY SCANNER",DATE],[""],["Metric","Value"],
            ["Candidates",str(len(cands))],["Multi-source",str(len([c for c in cands if c.sources>=2]))],
            ["Events",str(len(events))],[""],["SUMMARY"],[ex.get("summary","")],
            [""],["ACTIONS"],*[[a] for a in ex.get("actions",[])]],range_name="A1")
        ws2 = sh.add_worksheet("Opportunities",100,10)
        rows = [["#","Title","Category","Score","Sources","Urgency","Revenue Signal","Source URL"]]
        for i,o in enumerate(ai.get("opportunities",[])[:20],1):
            mc = _match_cand(o.get("title",""),cands)
            rows.append([str(i),o.get("title",""),o.get("category",""),
                str(mc.score) if mc else "-",str(mc.sources) if mc else "-",
                o.get("urgency",""),o.get("revenue_signal",""),mc.url if mc else ""])
        ws2.update(values=rows,range_name="A1")
        ws4 = sh.add_worksheet("History",200,5)
        hrows = [["Date","Title","Score","Sources","Category"]]
        for c in cands[:30]: hrows.append([DATE,c.title[:80],str(c.score),str(c.sources),c.category])
        ws4.update(values=hrows,range_name="A1")
        print(f"OK ({sh.url})"); return sh.url
    except Exception as e: log.error(f"Sheets: {e}"); return None

# =============================================================================
# SECTION 10 - HTML DASHBOARD
# =============================================================================

def build_html(cands, ai, events, all_sig, top3=None, trends=None):
    print("\nHTML DASHBOARD...", end=" ")
    ex = ai.get("executive",{}); opps = ai.get("opportunities",[])
    ground = ai.get("ground_intel",{}); comp_intel = ai.get("competitor_intel",{})
    top3 = top3 or {"GMG":[],"ENT":[],"PPM":[],"MTU":[]}
    trends = trends or {"kpi_deltas":{},"movers":[],"new_entries":[],"dropped":[]}
    total_sig = sum(len(v) for v in all_sig.values())
    multi = len([c for c in cands if c.sources>=2])
    n_cats = len(set(c.category for c in cands))
    crit_ev = len([e for e in events if e.get("urgency")=="critical"])
    n_sources = len([k for k,v in all_sig.items() if len(v)>0])
    urg_crit = len([o for o in opps if o.get("urgency")=="critical"])
    urg_high = len([o for o in opps if o.get("urgency")=="high"])
    kpi_deltas = trends.get("kpi_deltas", {})

    cat_count = defaultdict(int)
    for c in cands: cat_count[c.category] += 1
    top_cats = sorted(cat_count.items(),key=lambda x:-x[1])[:10]
    cat_labels = json.dumps([c[0] for c in top_cats]); cat_values = json.dumps([c[1] for c in top_cats])
    active_sources = {k:len(v) for k,v in all_sig.items() if len(v)>0}
    src_labels = json.dumps(list(active_sources.keys())); src_values = json.dumps(list(active_sources.values()))

    score_data = []
    for o in opps[:15]:
        mc = _match_cand(o.get("title",""),cands)
        score_data.append({"label":o.get("title","")[:40],"score":mc.score if mc else 0})
    score_labels = json.dumps([d["label"] for d in score_data])
    score_values = json.dumps([d["score"] for d in score_data])

    opp_rows = ""
    for i,o in enumerate(opps[:15],1):
        mc = _match_cand(o.get("title",""),cands)
        sc = f"{mc.score}" if mc else "-"; sr = str(mc.sources) if mc else "-"; url = mc.url if mc else ""
        urg = o.get("urgency",""); urg_cls = {"critical":"urg-crit","high":"urg-high","medium":"urg-med"}.get(urg,"urg-med")
        title_html = f'<a href="{esc(url)}" target="_blank" rel="noopener">{esc(o.get("title",""))}</a>' if url else esc(o.get("title",""))
        dom = _domain(url)
        src_html = f'<a href="{esc(url)}" target="_blank" class="src-link">{esc(dom)}</a>' if dom else '<span class="t2">-</span>'
        opp_rows += f"""<tr><td class="rank">{i}</td><td class="opp-title">{title_html}</td>
<td><span class="cat-tag">{esc(o.get('category',''))}</span></td><td class="score-val">{sc}</td><td>{sr}</td>
<td><span class="badge {urg_cls}">{urg.upper()}</span></td><td class="rev-sig">{esc(o.get('revenue_signal',''))[:80]}</td>
<td class="src-cell">{src_html}</td></tr>"""

    events_rows = ""
    for e in events[:20]:
        urg = e.get("urgency",""); urg_cls = {"critical":"urg-crit","high":"urg-high","medium":"urg-med"}.get(urg,"")
        live = ' <span class="badge badge-live">LIVE</span>' if e.get("is_live") else ""
        events_rows += f"""<tr><td>{esc(e['name'][:55])}{live}</td><td><span class="cat-tag">{esc(e['category'])}</span></td>
<td><span class="badge {urg_cls}">{esc(e['status'])}</span></td><td>{esc(e['description'][:50])}</td></tr>"""

    # Build Top 3 per Business Category HTML
    top3_html = ""
    for bc_code in ["GMG","ENT","PPM","MTU"]:
        bc_name = BIZ_CAT_NAMES[bc_code]; bc_emoji = BIZ_CAT_EMOJI[bc_code]
        bc_color = BIZ_CAT_COLORS[bc_code]; items = top3.get(bc_code, [])
        if not items:
            top3_html += f'<div class="biz-cat-section" style="border-left:3px solid {bc_color}"><h3>{bc_emoji} {esc(bc_name)} <span class="cat-tag" style="background:{bc_color}22;color:{bc_color}">{bc_code}</span></h3><p class="t2">No items this week.</p></div>'
            continue
        rows = ""
        for i, item in enumerate(items, 1):
            urg = item.get("urgency",""); urg_cls = {"critical":"urg-crit","high":"urg-high","medium":"urg-med"}.get(urg,"urg-med")
            url = item.get("url","")
            title_html = f'<a href="{esc(url)}" target="_blank">{esc(item["title"][:60])}</a>' if url else esc(item["title"][:60])
            rows += f'<tr><td class="rank">{i}</td><td class="opp-title">{title_html}</td><td class="score-val">{item.get("score",0)}</td><td><span class="badge {urg_cls}">{urg.upper()}</span></td><td class="rev-sig">{esc(item.get("why_now",""))[:80]}</td></tr>'
        top3_html += f'<div class="biz-cat-section" style="border-left:3px solid {bc_color}"><h3>{bc_emoji} {esc(bc_name)} <span class="cat-tag" style="background:{bc_color}22;color:{bc_color}">{bc_code}</span></h3><table><thead><tr><th>#</th><th>Story</th><th>Score</th><th>Urgency</th><th>Why Now</th></tr></thead><tbody>{rows}</tbody></table></div>'

    # Build advanced competitor intelligence HTML
    comp_intel_data = comp_intel if isinstance(comp_intel, dict) else {"competitors": comp_intel}
    comp_list = comp_intel_data.get("competitors", []) if isinstance(comp_intel_data, dict) else []
    market_summary = comp_intel_data.get("market_summary","") if isinstance(comp_intel_data, dict) else ""
    biggest_threat = comp_intel_data.get("biggest_threat","") if isinstance(comp_intel_data, dict) else ""
    opp_gap = comp_intel_data.get("opportunity_gap","") if isinstance(comp_intel_data, dict) else ""

    comp_intel_html = ""
    # Market overview bar
    if market_summary or biggest_threat or opp_gap:
        comp_intel_html += '<div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:10px;margin-bottom:16px">'
        if market_summary:
            comp_intel_html += f'<div style="background:var(--bg);border:1px solid var(--border);border-radius:8px;padding:12px"><div style="font-size:10px;color:var(--t2);text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px">Market Landscape</div><p style="font-size:12px;line-height:1.5">{esc(market_summary)}</p></div>'
        if biggest_threat:
            comp_intel_html += f'<div style="background:rgba(248,113,113,.05);border:1px solid rgba(248,113,113,.2);border-radius:8px;padding:12px"><div style="font-size:10px;color:var(--red);text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px">\U0001F6A8 Biggest Threat</div><p style="font-size:12px;line-height:1.5">{esc(biggest_threat)}</p></div>'
        if opp_gap:
            comp_intel_html += f'<div style="background:rgba(52,211,153,.05);border:1px solid rgba(52,211,153,.2);border-radius:8px;padding:12px"><div style="font-size:10px;color:var(--green);text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px">\U0001F4A1 Opportunity Gap</div><p style="font-size:12px;line-height:1.5">{esc(opp_gap)}</p></div>'
        comp_intel_html += '</div>'

    # Per-competitor cards
    if comp_list:
        for ci in comp_list:
            cn = ci.get("name",""); tl = ci.get("threat_level","medium")
            ct = ci.get("type","direct")
            tl_cls = {"high":"threat-high","medium":"threat-medium","low":"threat-low"}.get(tl,"threat-medium")
            tl_bg = {"high":"rgba(248,113,113,.08)","medium":"rgba(251,191,36,.08)","low":"rgba(52,211,153,.08)"}.get(tl,"rgba(251,191,36,.08)")
            type_badge = f'<span class="pill" style="margin-left:6px">{ct.upper()}</span>' if ct else ""
            acts = "".join(f'<li>{esc(a)}</li>' for a in ci.get("activities",[])[:5])
            news = "".join(f'<li>{esc(n)}</li>' for n in ci.get("news_highlights",[])[:3])
            notable = ci.get("notable","")
            strategy = ci.get("strategy_signal","")
            response = ci.get("our_response","")
            web_moves = ci.get("website_moves","")

            comp_intel_html += f'<div class="comp-intel-card" style="background:{tl_bg};border-color:var(--border)">'
            comp_intel_html += f'<h4>{esc(cn)}{type_badge} <span class="{tl_cls}" style="font-size:10px;text-transform:uppercase;letter-spacing:.3px">{tl.upper()} THREAT</span></h4>'
            comp_intel_html += f'<p style="font-size:13px;font-weight:500;margin-bottom:8px">{esc(ci.get("headline",""))}</p>'
            if strategy:
                comp_intel_html += f'<p style="font-size:11.5px;color:var(--purple);margin-bottom:8px"><strong>Strategy Signal:</strong> {esc(strategy)}</p>'
            if acts:
                comp_intel_html += f'<div style="margin-bottom:8px"><div style="font-size:10px;color:var(--t2);text-transform:uppercase;margin-bottom:4px">Activities</div><ul style="font-size:12px;padding-left:16px;margin:0">{acts}</ul></div>'
            if news:
                comp_intel_html += f'<div style="margin-bottom:8px"><div style="font-size:10px;color:var(--blue);text-transform:uppercase;margin-bottom:4px">In The News</div><ul style="font-size:12px;padding-left:16px;margin:0">{news}</ul></div>'
            if web_moves:
                comp_intel_html += f'<p style="font-size:11.5px;margin-bottom:6px"><strong style="color:var(--t2)">Website:</strong> {esc(web_moves)}</p>'
            if notable:
                comp_intel_html += f'<p style="font-size:12px;color:var(--amber);margin-bottom:6px">\U0001F50E <strong>Notable:</strong> {esc(notable)}</p>'
            if response:
                comp_intel_html += f'<p style="font-size:12px;color:var(--green);background:rgba(52,211,153,.08);padding:6px 10px;border-radius:6px;margin-top:6px">\U0001F3AF <strong>Our Move:</strong> {esc(response)}</p>'
            comp_intel_html += '</div>'
    else:
        comp_by = defaultdict(list)
        for s in all_sig.get("competitor",[]):
            p = s.meta.get("product","")
            if p and p not in comp_by[s.meta.get("comp","?")]: comp_by[s.meta.get("comp","?")].append(p)
        comp_intel_html += "".join(f'<div class="comp-row"><strong>{esc(cn)}</strong>{"".join(f"""<span class="pill">{esc(p)}</span>""" for p in prods[:8])}</div>' for cn,prods in comp_by.items())
    comp_html = comp_intel_html

    # Verified Sources section removed per user request

    pred_html = "".join(f"<li>{esc(p)}</li>" for p in ex.get("predictions",[]))
    risk_html = "".join(f"<li>{esc(r)}</li>" for r in ex.get("risks",[]))

    # Pre-compute hero category highlights (avoids nested f-string {{}} issue)
    cat_highlights = ex.get("category_highlights", {})
    hero_highlights_html = ""
    highlight_divs = []
    for bc in ["GMG","ENT","PPM","MTU"]:
        hl_text = cat_highlights.get(bc, "")
        if hl_text:
            highlight_divs.append(
                f'<div style="padding:8px 12px;background:{BIZ_CAT_COLORS[bc]}11;'
                f'border-left:3px solid {BIZ_CAT_COLORS[bc]};border-radius:0 6px 6px 0;'
                f'font-size:12px"><strong>{BIZ_CAT_EMOJI[bc]} {BIZ_CAT_NAMES[bc]}:</strong> '
                f'{esc(hl_text)}</div>'
            )
    if highlight_divs:
        hero_highlights_html = '<div style="display:grid;grid-template-columns:1fr 1fr;gap:8px;margin-top:12px">' + "".join(highlight_divs) + '</div>'

    # Action Items section removed per user request

    # Build per-competitor sitemap + blog sections
    sitemap_sigs = all_sig.get("sitemap",[])
    sitemap_by_comp = defaultdict(list)
    for s in sitemap_sigs:
        if s.meta.get("type") == "blog":
            sitemap_by_comp[s.meta.get("comp","?")].append(s)
    sitemap_html = ""
    for comp_name in SITEMAP_COMPETITORS:
        blogs = sitemap_by_comp.get(comp_name,[])
        if blogs:
            rows = ""
            for s in blogs[:15]:
                page_title = s.title.replace(f"{comp_name}: ","",1)
                lastmod = s.meta.get("lastmod","")
                rows += f"""<tr><td class="opp-title"><a href="{esc(s.url)}" target="_blank" rel="noopener">{esc(page_title[:60])}</a></td><td>{esc(lastmod)}</td></tr>"""
            sitemap_html += f"""<h3 style="margin-top:14px">{esc(comp_name)} <span class="badge urg-high">{len(blogs)} blog post{'s' if len(blogs)!=1 else ''}</span></h3>
<table><thead><tr><th>Title</th><th>Date</th></tr></thead><tbody>{rows}</tbody></table>"""
        else:
            sitemap_html += f"""<h3 style="margin-top:14px">{esc(comp_name)} <span class="t2" style="font-weight:normal;font-size:11px">&mdash; no new blog posts this week</span></h3>"""
    if not sitemap_html:
        sitemap_html = '<p class="t2">No competitor blog posts found.</p>'

    # Build KPI delta HTML
    def _delta_html(key):
        d = kpi_deltas.get(key, 0)
        if d > 0: return f'<div class="delta up">+{d} vs last week</div>'
        elif d < 0: return f'<div class="delta down">{d} vs last week</div>'
        return '<div class="delta neutral">no change</div>'

    html = f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Recharge.com Opportunity Scanner</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.7/dist/chart.umd.min.js"></script>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
:root{{
  --bg:#0e1b1c;--bg2:#142222;--card:#192829;--card2:#1e3030;--elevated:#223636;
  --border:#2a4040;--border-h:#356060;--border-a:rgba(103,223,136,.3);
  --t:#f0f1f5;--t2:#8b9a9c;--t3:#5c7070;--t-inv:#0e1b1c;
  --accent:#67df88;--accent-h:#52c970;--accent-s:rgba(103,223,136,.08);--accent-g:rgba(103,223,136,.15);
  --green:#67df88;--green-s:rgba(103,223,136,.08);--red:#f87171;--red-s:rgba(248,113,113,.08);
  --amber:#fbbf24;--amber-s:rgba(251,191,36,.08);--blue:#60a5fa;--blue-s:rgba(96,165,250,.08);--purple:#c084fc;
  --font:'Inter',-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;
  --sp1:4px;--sp2:8px;--sp3:12px;--sp4:16px;--sp5:20px;--sp6:24px;--sp8:32px;
  --r-sm:6px;--r-md:10px;--r-lg:14px;--r-xl:20px;
  --sh-sm:0 1px 2px rgba(0,0,0,.3);--sh-md:0 4px 12px rgba(0,0,0,.25);--sh-lg:0 8px 30px rgba(0,0,0,.35);--sh-glow:0 0 20px rgba(129,140,248,.1);
  --ease:cubic-bezier(.4,0,.2,1);--dur:200ms
}}
body{{font-family:var(--font);background:var(--bg);color:var(--t);line-height:1.6;font-size:14px;-webkit-font-smoothing:antialiased;-moz-osx-font-smoothing:grayscale}}
a{{color:var(--accent);text-decoration:none}}a:hover{{color:var(--accent-h)}}
:focus-visible{{outline:2px solid var(--accent);outline-offset:2px;border-radius:4px}}
.skip-link{{position:absolute;top:-40px;left:0;background:var(--accent);color:#fff;padding:8px 16px;z-index:100;transition:top .2s}}.skip-link:focus{{top:0}}
.wrap{{max-width:1440px;margin:0 auto;padding:var(--sp5) var(--sp6)}}
header{{display:flex;justify-content:space-between;align-items:flex-end;margin-bottom:var(--sp8);padding-bottom:var(--sp4);border-bottom:1px solid var(--border)}}
header h1{{font-size:24px;font-weight:800;letter-spacing:-.025em}}header h1 span{{color:var(--accent)}}
.meta{{color:var(--t2);font-size:12px;text-align:right}}.meta strong{{color:var(--green)}}
.kpis{{display:grid;grid-template-columns:repeat(auto-fit,minmax(160px,1fr));gap:var(--sp3);margin-bottom:var(--sp6)}}
.kpi{{background:var(--card);border:1px solid var(--border);border-radius:var(--r-lg);padding:var(--sp5) var(--sp4);transition:all var(--dur) var(--ease);position:relative;overflow:hidden}}
.kpi:hover{{border-color:var(--border-h);transform:translateY(-1px);box-shadow:var(--sh-md)}}
.kpi::before{{content:'';position:absolute;top:0;left:0;right:0;height:2px;background:var(--accent);opacity:0.5;transition:opacity var(--dur)}}.kpi:hover::before{{opacity:1}}
.kpi .lb{{font-size:11px;font-weight:600;color:var(--t3);text-transform:uppercase;letter-spacing:.06em;margin-bottom:var(--sp1)}}
.kpi .vl{{font-size:28px;font-weight:800;font-variant-numeric:tabular-nums;line-height:1.1}}
.kpi .vl.green{{color:var(--green)}}.kpi .vl.amber{{color:var(--amber)}}.kpi .vl.red{{color:var(--red)}}
.kpi .vl.blue{{color:var(--blue)}}.kpi .vl.purple{{color:var(--purple)}}
.delta{{font-size:11px;font-weight:600;margin-top:var(--sp1);display:flex;align-items:center;gap:3px}}
.delta.up{{color:var(--green)}}.delta.down{{color:var(--red)}}.delta.neutral{{color:var(--t3)}}
@keyframes fadeUp{{from{{opacity:0;transform:translateY(12px)}}to{{opacity:1;transform:translateY(0)}}}}
.card{{background:var(--card);border:1px solid var(--border);border-radius:var(--r-lg);padding:var(--sp5);margin-bottom:var(--sp5);animation:fadeUp .4s var(--ease) backwards;box-shadow:var(--sh-sm)}}
.card:nth-child(2){{animation-delay:.05s}}.card:nth-child(3){{animation-delay:.1s}}.card:nth-child(4){{animation-delay:.15s}}.card:nth-child(5){{animation-delay:.2s}}
.card h2{{font-size:16px;font-weight:700;margin-bottom:var(--sp4);padding-bottom:var(--sp2);border-bottom:1px solid var(--border);letter-spacing:-.01em;padding-left:var(--sp3);border-left:3px solid var(--accent)}}
.card h3{{font-size:13px;font-weight:600;margin:var(--sp4) 0 var(--sp2);color:var(--t2)}}
.collapsible{{cursor:pointer;user-select:none;display:flex;justify-content:space-between;align-items:center}}
.collapsible::after{{content:'';width:8px;height:8px;border-right:2px solid var(--t3);border-bottom:2px solid var(--t3);transform:rotate(45deg);transition:transform var(--dur) var(--ease);flex-shrink:0}}
.collapsible.collapsed::after{{transform:rotate(-45deg)}}
.collapsible-content{{max-height:5000px;overflow:hidden;transition:max-height .4s var(--ease),opacity .3s;opacity:1}}
.collapsible-content.hidden{{max-height:0;opacity:0}}
.summary-text{{font-size:15px;line-height:1.7}}
.action-item{{padding:8px 0;border-bottom:1px solid var(--border);font-size:13px;transition:background var(--dur)}}.action-item:last-child{{border:none}}.action-item:hover{{background:var(--accent-s);padding-left:8px;border-radius:var(--r-sm)}}
.table-wrap{{overflow-x:auto;-webkit-overflow-scrolling:touch}}
table{{width:100%;border-collapse:collapse;font-size:13px}}
th{{text-align:left;padding:10px 12px;background:var(--card2);color:var(--t2);font-weight:600;text-transform:uppercase;font-size:10.5px;letter-spacing:.5px;border-bottom:2px solid var(--border);cursor:pointer;user-select:none;transition:color var(--dur)}}
th:hover{{color:var(--accent)}}th[data-sort]::after{{content:' \u2195';font-size:9px;opacity:.4}}
td{{padding:10px 12px;border-bottom:1px solid var(--border)}}tr{{transition:background var(--dur) var(--ease)}}tr:nth-child(even){{background:rgba(129,140,248,.03)}}tr:hover{{background:var(--accent-s)}}
.rank{{color:var(--t2);font-weight:600;width:30px}}.opp-title{{font-weight:600;max-width:300px}}
.opp-title a{{color:var(--t);text-decoration:none;border-bottom:1px dotted var(--t3);transition:all .15s}}
.opp-title a:hover{{color:var(--accent);border-color:var(--accent)}}
.score-val{{font-weight:700;color:var(--accent);font-size:15px;font-variant-numeric:tabular-nums}}
.rev-sig{{color:var(--t2);font-size:12px;max-width:260px}}.src-cell{{font-size:12px}}
.src-link{{color:var(--blue);text-decoration:none;font-size:11px}}.src-link:hover{{text-decoration:underline}}
.t2{{color:var(--t2)}}
.cat-tag{{background:rgba(129,140,248,.1);color:var(--accent);padding:2px 8px;border-radius:var(--r-sm);font-size:11px;white-space:nowrap}}
.badge{{display:inline-block;padding:4px 12px;border-radius:20px;font-size:10.5px;font-weight:600;text-transform:uppercase;letter-spacing:.3px}}
.urg-crit{{background:var(--red-s);color:var(--red)}}.urg-high{{background:var(--amber-s);color:var(--amber)}}
.urg-med{{background:var(--accent-s);color:var(--accent)}}
.badge-live{{background:rgba(248,113,113,.18);color:var(--red);animation:livePulse 2s ease-in-out infinite}}
@keyframes livePulse{{0%,100%{{box-shadow:0 0 0 0 rgba(248,113,113,.4)}}50%{{box-shadow:0 0 0 6px rgba(248,113,113,0)}}}}
.chart-row{{display:grid;gap:var(--sp4);margin-bottom:var(--sp5)}}.chart-hero{{grid-template-columns:1fr}}.chart-pair{{grid-template-columns:1fr 1fr}}
.chart-box{{background:var(--card);border:1px solid var(--border);border-radius:var(--r-lg);padding:18px;box-shadow:var(--sh-sm),inset 0 1px 0 rgba(129,140,248,.06)}}
.chart-box h3{{font-size:12px;color:var(--t2);margin-bottom:10px;text-transform:uppercase;letter-spacing:.5px;border:none;padding:0}}
.chart-container{{position:relative;height:380px}}
.comp-intel-card ul{{list-style:disc}}.comp-intel-card li{{margin-bottom:3px}}
.comp-row{{display:flex;align-items:center;gap:6px;flex-wrap:wrap;margin-bottom:6px}}
.pill{{background:var(--accent-s);color:var(--accent);padding:2px 10px;border-radius:12px;font-size:11px}}
.ground-sources{{margin-top:12px;padding-top:10px;border-top:1px solid var(--border)}}
.ground-sources h4{{font-size:11px;color:var(--t2);text-transform:uppercase;letter-spacing:.5px;margin-bottom:6px}}
.ground-sources ul{{list-style:none;padding:0}}.ground-sources li{{font-size:12px;margin-bottom:3px}}
.ground-sources a{{color:var(--blue);text-decoration:none}}.ground-sources a:hover{{text-decoration:underline}}
.two-col{{display:grid;grid-template-columns:1fr 1fr;gap:var(--sp5)}}.two-col .card{{margin-bottom:0}}
.biz-cat-section{{margin-bottom:var(--sp4);padding:var(--sp4);background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-md);transition:border-color var(--dur)}}.biz-cat-section:hover{{border-color:var(--border-h)}}
.biz-cat-section h3{{font-size:14px;font-weight:600;margin-bottom:10px;display:flex;align-items:center;gap:8px}}
.biz-cat-grid{{display:grid;grid-template-columns:1fr 1fr;gap:var(--sp4)}}
.comp-intel-card{{background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-md);padding:var(--sp4);margin-bottom:10px;transition:border-color var(--dur)}}.comp-intel-card:hover{{border-color:var(--border-h)}}
.comp-intel-card h4{{font-size:13px;font-weight:600;margin-bottom:6px;display:flex;justify-content:space-between;align-items:center}}
.threat-indicator{{display:inline-block;width:8px;height:8px;border-radius:50%;margin-right:6px}}
.threat-indicator.threat-high{{background:var(--red);box-shadow:0 0 6px var(--red)}}.threat-indicator.threat-medium{{background:var(--amber)}}.threat-indicator.threat-low{{background:var(--green)}}
.threat-high{{color:var(--red)}}.threat-medium{{color:var(--amber)}}.threat-low{{color:var(--green)}}
.search-input{{background:var(--bg);border:1px solid var(--border);border-radius:var(--r-sm);padding:8px 14px;color:var(--t);font-size:13px;width:260px;outline:none;transition:border-color var(--dur);font-family:var(--font)}}.search-input:focus{{border-color:var(--accent)}}
::-webkit-scrollbar{{width:6px;height:6px}}::-webkit-scrollbar-track{{background:var(--bg)}}::-webkit-scrollbar-thumb{{background:var(--border);border-radius:3px}}::-webkit-scrollbar-thumb:hover{{background:var(--t3)}}
@media print{{body{{background:#fff;color:#111;font-size:12px}}.wrap{{max-width:100%;padding:0}}.card,.kpi{{border:1px solid #ddd;background:#fff;box-shadow:none;break-inside:avoid}}.kpi .vl{{color:#111}}.badge{{border:1px solid #999;background:#f0f0f0;color:#333}}.urg-crit{{border-color:#e00;color:#c00}}.urg-high{{border-color:#c80;color:#a60}}header{{border-bottom:2px solid #111}}a{{color:#111;text-decoration:underline}}.chart-box,canvas{{display:none}}.table-wrap{{overflow:visible}}tr:hover{{background:none}}@page{{margin:1.5cm}}.search-input{{display:none}}}}
@media(max-width:1200px){{.kpis{{grid-template-columns:repeat(4,1fr)}}.chart-pair{{grid-template-columns:1fr}}}}
@media(max-width:900px){{.wrap{{padding:12px 16px}}.kpis{{grid-template-columns:repeat(3,1fr)}}.two-col,.biz-cat-grid,.chart-pair{{grid-template-columns:1fr}}header{{flex-direction:column;align-items:flex-start;gap:8px}}.meta{{text-align:left}}.kpi .vl{{font-size:22px}}}}
@media(max-width:600px){{.kpis{{grid-template-columns:repeat(2,1fr);gap:8px}}.kpi{{padding:12px}}.kpi .vl{{font-size:20px}}.kpi .lb{{font-size:10px}}table{{font-size:12px}}td,th{{padding:6px 8px}}.opp-title{{max-width:140px}}.rev-sig,.src-cell{{display:none}}.card{{padding:14px;border-radius:var(--r-md)}}h1{{font-size:18px}}.search-input{{width:100%}}}}
</style></head><body>
<a href="#main" class="skip-link">Skip to main content</a>
<div class="wrap">
<header role="banner"><div style="display:flex;align-items:center;gap:14px"><img src="https://company.recharge.com/wp-content/uploads/2024/10/Recharge.com-Logo.svg" alt="Recharge.com" style="height:32px;filter:brightness(0) invert(1)" /><h1>Opportunity Scanner</h1></div>
<div class="meta">v5.0 &middot; {DATE} {TIME}<br><strong>{n_sources} sources</strong> &middot; 4-pass AI (grounded)</div></header>
<main role="main" id="main">
<section aria-label="Key Performance Indicators">
<div class="kpis">
<div class="kpi"><div class="lb">Signals</div><div class="vl blue">{total_sig}</div>{_delta_html('total_signals')}</div>
<div class="kpi"><div class="lb">Candidates</div><div class="vl">{len(cands)}</div>{_delta_html('candidates')}</div>
<div class="kpi"><div class="lb">Multi-Source</div><div class="vl purple">{multi}</div>{_delta_html('multi_source')}</div>
<div class="kpi"><div class="lb">Categories</div><div class="vl green">{n_cats}</div></div>
<div class="kpi"><div class="lb">Critical Events</div><div class="vl red">{crit_ev}</div></div>
<div class="kpi"><div class="lb">Critical Opps</div><div class="vl red">{urg_crit}</div></div>
<div class="kpi"><div class="lb">High Priority</div><div class="vl amber">{urg_high}</div></div>
</div></section>
<div class="card" style="background:linear-gradient(135deg,var(--card),var(--card2),rgba(129,140,248,.06));border:1px solid var(--accent);box-shadow:0 0 30px rgba(129,140,248,.08)">
<div style="display:flex;gap:20px;align-items:flex-start">
<div style="flex:1">
<h2 style="border:none;margin-bottom:8px;color:var(--accent);font-size:16px">This Week's Intelligence Brief</h2>
<p class="summary-text" style="margin-bottom:12px">{esc(ex.get('summary',''))}</p>
{hero_highlights_html}
</div>
<div style="min-width:120px;text-align:right">
<div style="font-size:42px;font-weight:800;color:var(--accent)">{urg_crit + urg_high}</div>
<div style="font-size:11px;color:var(--t2);text-transform:uppercase">Priority Items</div>
</div></div></div>
<div class="chart-row chart-hero"><div class="chart-box"><h3>Top Opportunities by Composite Score</h3>
<div class="chart-container"><canvas id="scoreChart"></canvas></div></div></div>
<div class="chart-row chart-pair">
<div class="chart-box"><h3>Signals by Source</h3><canvas id="srcChart"></canvas></div>
<div class="chart-box"><h3>Candidates by Category</h3><canvas id="catChart"></canvas></div></div>
<div class="card"><div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:12px;flex-wrap:wrap;gap:8px">
<h2 style="border:none;margin:0;padding:0">Top Opportunities</h2>
<input id="oppSearch" type="text" placeholder="Search opportunities..." class="search-input"
 onfocus="this.style.borderColor='var(--accent)'" onblur="this.style.borderColor='var(--border)'" />
</div><div class="table-wrap"><table id="oppTable" role="grid" aria-label="Top Opportunities"><thead><tr>
<th scope="col" data-sort="num">#</th><th scope="col" data-sort="text">Opportunity</th><th scope="col" data-sort="text">Category</th><th scope="col" data-sort="num">Score</th><th scope="col" data-sort="num">Sources</th><th scope="col" data-sort="text">Urgency</th><th scope="col">Revenue Signal</th><th scope="col">Source</th>
</tr></thead><tbody>{opp_rows}</tbody></table></div></div>
<div class="card"><h2>\U0001F4CA This Week's Must-Know: Top 3 Per Category</h2><div class="biz-cat-grid">{top3_html}</div></div>
<div class="card"><h2>\U0001F50D Competitor Intelligence Center</h2>{comp_html if comp_html else '<p class="t2">No competitor data.</p>'}</div>
<div class="card"><h2>\U0001F4C4 Competitor Blog Activity</h2>{sitemap_html}</div>
<div class="two-col">
<div class="card"><h2>\U0001F52E Predicted Trends</h2><ul style="font-size:13px;padding-left:16px">{pred_html}</ul></div>
<div class="card"><h2>\U0001F6A8 Risk Watchlist</h2><ul style="font-size:13px;padding-left:16px">{risk_html}</ul></div></div>
<div class="card"><h2 class="collapsible">Events Calendar</h2><div class="collapsible-content"><div class="table-wrap"><table><thead><tr><th scope="col">Event</th><th scope="col">Category</th><th scope="col">Status</th><th scope="col">Details</th></tr></thead>
<tbody>{events_rows}</tbody></table></div></div></div>
<div class="card"><h2 class="collapsible collapsed">How It Works</h2><div class="collapsible-content hidden">
<div style="margin-bottom:20px">
<h3 style="color:var(--accent);font-size:14px;margin-bottom:8px">What This Tool Does</h3>
<p style="font-size:13px;line-height:1.7;color:var(--t2)">The Recharge.com Opportunity Scanner automatically scans {n_sources}+ data sources across the internet every week to find sales opportunities for our business. It collects signals from gaming platforms, news sites, competitor websites, social media, and deal aggregators, then uses AI to rank and prioritize the most actionable opportunities for our team.</p>
</div>
<div style="margin-bottom:20px">
<h3 style="color:var(--accent);font-size:14px;margin-bottom:8px">Data Sources ({n_sources} active)</h3>
<div style="display:grid;grid-template-columns:1fr 1fr;gap:8px">
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">Steam</strong><br><span class="t2">Featured games, top sellers, sales, new releases</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">Reddit</strong><br><span class="t2">{len(SUBREDDITS)} gaming subreddits — what gamers are discussing</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">News RSS</strong><br><span class="t2">{len(RSS_FEEDS)} news sites (IGN, GameSpot, PYMNTS, etc.)</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">Google News</strong><br><span class="t2">Real-time news search via Oxylabs across all categories</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">Competitors</strong><br><span class="t2">Homepage scanning + blog monitoring for {', '.join(COMPETITORS.keys())}</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">CheapShark</strong><br><span class="t2">PC game deal aggregator — best prices across stores</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">SteamSpy</strong><br><span class="t2">Player counts and trending games data</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">Epic / GOG / Humble</strong><br><span class="t2">Free games, sales, and bestsellers from each store</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">GamerPower</strong><br><span class="t2">Free game giveaways across all platforms</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">YouTube</strong><br><span class="t2">{len(YT_CHANNELS)} gaming channels — latest video uploads</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">Anime (Jikan)</strong><br><span class="t2">Top airing + upcoming anime for Crunchyroll insights</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--accent)">Wikipedia / Trends</strong><br><span class="t2">Page view data + Google search interest trends</span></div>
</div></div>
<div style="margin-bottom:20px">
<h3 style="color:var(--accent);font-size:14px;margin-bottom:8px">How We Categorize</h3>
<p style="font-size:13px;line-height:1.7;color:var(--t2);margin-bottom:8px">Every signal is matched against {len(KW)} keyword categories using text matching. These categories roll up into 4 business segments:</p>
<div style="display:flex;gap:8px;flex-wrap:wrap">
<span style="background:{BIZ_CAT_COLORS['GMG']}22;color:{BIZ_CAT_COLORS['GMG']};padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600">{BIZ_CAT_EMOJI['GMG']} GMG — Gaming</span>
<span style="background:{BIZ_CAT_COLORS['ENT']}22;color:{BIZ_CAT_COLORS['ENT']};padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600">{BIZ_CAT_EMOJI['ENT']} ENT — Entertainment</span>
<span style="background:{BIZ_CAT_COLORS['PPM']}22;color:{BIZ_CAT_COLORS['PPM']};padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600">{BIZ_CAT_EMOJI['PPM']} PPM — Prepaid Money</span>
<span style="background:{BIZ_CAT_COLORS['MTU']}22;color:{BIZ_CAT_COLORS['MTU']};padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600">{BIZ_CAT_EMOJI['MTU']} MTU — Mobile Top Up</span>
</div></div>
<div style="margin-bottom:20px">
<h3 style="color:var(--accent);font-size:14px;margin-bottom:8px">How Scoring Works</h3>
<p style="font-size:13px;line-height:1.7;color:var(--t2)">Each source's raw scores are <strong>normalized to 0-100</strong> using per-source floor/ceiling ranges so they can be compared fairly. Then a <strong>weighted composite score</strong> is calculated: Tier 1 sources (Oxylabs, Trends, News, Reddit) carry ~50% of the weight, Tier 2 (Steam, Epic, YouTube, Competitors) ~30%, and Tier 3 (deal sites, anime, wiki) ~20%. Items appearing in <strong>multiple sources</strong> get a confidence multiplier (up to 1.0x for 4+ sources), and a <strong>freshness factor</strong> boosts very recent items.</p>
</div>
<div>
<h3 style="color:var(--accent);font-size:14px;margin-bottom:8px">How AI Analysis Works</h3>
<p style="font-size:13px;line-height:1.7;color:var(--t2)">Google Gemini AI reviews all the data in a <strong>4-pass pipeline</strong>:</p>
<div style="display:grid;grid-template-columns:1fr 1fr;gap:8px;margin-top:8px">
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--green)">Pass 1: Fact-Check</strong><br><span class="t2">Verifies recent developments via Google Search grounding</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--amber)">Pass 2: Prioritize</strong><br><span class="t2">Picks the top 15 opportunities based on timing and revenue impact</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--blue)">Pass 3: Executive Brief</strong><br><span class="t2">Writes a summary with actions, predictions, and risks</span></div>
<div style="background:var(--bg2);border:1px solid var(--border);border-radius:var(--r-sm);padding:10px 12px;font-size:12px"><strong style="color:var(--purple)">Pass 4: Competitor Intel</strong><br><span class="t2">Deep analysis of competitor strategies and threat levels</span></div>
</div></div>
</div></div>
</main>
<footer role="contentinfo" style="text-align:center;padding:20px;color:var(--t2);font-size:11px">
Recharge.com v5.0 &middot; {n_sources} sources &middot; {total_sig} signals &middot; Gemini (grounded)</footer></div>
<script>
const C=['#818cf8','#60a5fa','#34d399','#fbbf24','#f87171','#c084fc','#fb923c','#2dd4bf','#a78bfa','#f472b6','#38bdf8','#4ade80','#e879f9','#67e8f9','#bef264'];
const TT={{backgroundColor:'#1c2030',titleFont:{{family:'Inter',size:12,weight:600}},bodyFont:{{family:'Inter',size:11}},padding:12,borderColor:'#2a3050',borderWidth:1,cornerRadius:8,displayColors:true,boxPadding:4}};
Chart.defaults.color='#8b90a8';Chart.defaults.borderColor='#1e2235';Chart.defaults.font.family='Inter';
const sCtx=document.getElementById('scoreChart').getContext('2d');
new Chart(sCtx,{{type:'bar',
data:{{labels:{score_labels},datasets:[{{data:{score_values},
backgroundColor:{score_values}.map((v,i)=>{{const g=sCtx.createLinearGradient(0,0,400,0);g.addColorStop(0,'hsla('+(235+i*8)+',80%,70%,0.85)');g.addColorStop(1,'hsla('+(235+i*8)+',80%,70%,0.45)');return g}}),
borderRadius:{{topLeft:4,topRight:4,bottomLeft:4,bottomRight:4}},borderSkipped:false,barPercentage:0.6}}]}},
options:{{indexAxis:'y',responsive:true,maintainAspectRatio:false,animation:{{duration:800,easing:'easeOutQuart'}},
plugins:{{legend:{{display:false}},tooltip:TT}},scales:{{x:{{grid:{{color:'#1e2235',drawBorder:false}},max:100,ticks:{{font:{{size:10}}}}}},y:{{grid:{{display:false}},ticks:{{font:{{size:11,weight:600}},padding:8}}}}}}
}}}});
new Chart(document.getElementById('srcChart'),{{type:'doughnut',
data:{{labels:{src_labels},datasets:[{{data:{src_values},backgroundColor:C,borderWidth:0,hoverOffset:6}}]}},
options:{{responsive:true,animation:{{duration:800,easing:'easeOutQuart'}},plugins:{{legend:{{position:'right',labels:{{boxWidth:10,padding:8,font:{{size:11}},usePointStyle:true,pointStyle:'circle'}}}},tooltip:TT}},cutout:'65%'}}}});
new Chart(document.getElementById('catChart'),{{type:'bar',
data:{{labels:{cat_labels},datasets:[{{data:{cat_values},backgroundColor:'rgba(129,140,248,.7)',hoverBackgroundColor:'#818cf8',borderRadius:5,borderSkipped:false}}]}},
options:{{responsive:true,animation:{{duration:800,easing:'easeOutQuart'}},plugins:{{legend:{{display:false}},tooltip:TT}},scales:{{x:{{grid:{{display:false}},ticks:{{font:{{size:10}},maxRotation:45}}}},y:{{grid:{{color:'#1e2235',drawBorder:false}}}}}}}}}});

// Table sorting
document.querySelectorAll('th[data-sort]').forEach(th=>{{
  th.addEventListener('click',()=>{{
    const table=th.closest('table'),tbody=table.querySelector('tbody'),rows=Array.from(tbody.querySelectorAll('tr')),col=th.cellIndex,type=th.dataset.sort,asc=th.classList.toggle('sort-asc');
    table.querySelectorAll('th').forEach(h=>{{if(h!==th)h.classList.remove('sort-asc')}});
    rows.sort((a,b)=>{{let va=a.cells[col].textContent.trim(),vb=b.cells[col].textContent.trim();if(type==='num'){{va=parseFloat(va)||0;vb=parseFloat(vb)||0;return asc?va-vb:vb-va}}return asc?va.localeCompare(vb):vb.localeCompare(va)}});
    rows.forEach(r=>tbody.appendChild(r));
  }});
}});

// Search filter
const si=document.getElementById('oppSearch');
if(si){{si.addEventListener('input',e=>{{const q=e.target.value.toLowerCase();document.querySelectorAll('#oppTable tbody tr').forEach(row=>{{row.style.display=row.textContent.toLowerCase().includes(q)?'':'none'}})}})}};

// Collapsible sections
document.querySelectorAll('.collapsible').forEach(el=>{{
  el.addEventListener('click',()=>{{el.classList.toggle('collapsed');const c=el.nextElementSibling;if(c)c.classList.toggle('hidden')}});
}});
</script></body></html>"""
    # ---- password protection wrapper (AES-GCM or XOR fallback) ----
    if DASH_PASSWORD:
        import hashlib, base64, os as _os
        salt = _os.urandom(16)
        iv = _os.urandom(12)
        salt_b64 = base64.b64encode(salt).decode()
        iv_b64 = base64.b64encode(iv).decode()
        html_bytes = html.encode("utf-8")
        if HAS_CRYPTO:
            key = hashlib.pbkdf2_hmac('sha256', DASH_PASSWORD.encode(), salt, 100000)
            aesgcm = AESGCM(key)
            encrypted = aesgcm.encrypt(iv, html_bytes, None)
            enc_b64 = base64.b64encode(encrypted).decode()
            decrypt_js = f"""async function unlock(){{
  const pw=document.getElementById('pwInput').value;
  const enc=new TextEncoder();
  try{{
    const km=await crypto.subtle.importKey('raw',enc.encode(pw),'PBKDF2',false,['deriveKey']);
    const key=await crypto.subtle.deriveKey(
      {{name:'PBKDF2',salt:Uint8Array.from(atob('{salt_b64}'),c=>c.charCodeAt(0)),iterations:100000,hash:'SHA-256'}},
      km,{{name:'AES-GCM',length:256}},false,['decrypt']);
    const dec=await crypto.subtle.decrypt(
      {{name:'AES-GCM',iv:Uint8Array.from(atob('{iv_b64}'),c=>c.charCodeAt(0))}},
      key,Uint8Array.from(atob(D),c=>c.charCodeAt(0)));
    document.open();document.write(new TextDecoder().decode(dec));document.close();
  }}catch(e){{document.getElementById('pwErr').style.display='block'}}
}}"""
        else:
            pw_bytes = DASH_PASSWORD.encode()
            encrypted = bytes(b ^ pw_bytes[i % len(pw_bytes)] for i, b in enumerate(html_bytes))
            enc_b64 = base64.b64encode(encrypted).decode()
            pw_hash = hashlib.sha256(DASH_PASSWORD.encode()).hexdigest()
            decrypt_js = f"""async function unlock(){{
  const pw=document.getElementById('pwInput').value;
  const e=new TextEncoder().encode(pw);const h=await crypto.subtle.digest('SHA-256',e);
  const hx=Array.from(new Uint8Array(h)).map(b=>b.toString(16).padStart(2,'0')).join('');
  if(hx!=='{pw_hash}'){{document.getElementById('pwErr').style.display='block';return}}
  const enc=Uint8Array.from(atob(D),c=>c.charCodeAt(0));const pwB=new TextEncoder().encode(pw);
  const dec=new Uint8Array(enc.length);for(let i=0;i<enc.length;i++)dec[i]=enc[i]^pwB[i%pwB.length];
  document.open();document.write(new TextDecoder().decode(dec));document.close()
}}"""
        html = f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Recharge.com Dashboard - Login</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{background:#09090b;font-family:'Inter',-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;color:#f0f1f5;display:flex;align-items:center;justify-content:center;min-height:100vh;-webkit-font-smoothing:antialiased}}
.login-box{{background:#141620;border:1px solid #1e2235;border-radius:20px;padding:48px 40px;width:400px;text-align:center;box-shadow:0 8px 30px rgba(0,0,0,.35)}}
.login-box h1{{font-size:22px;font-weight:800;margin-bottom:6px;letter-spacing:-.02em}}.login-box h1 span{{color:#818cf8}}
.login-box p{{color:#8b90a8;font-size:13px;margin-bottom:28px}}
.pw-input{{width:100%;padding:14px 18px;background:#09090b;border:1px solid #1e2235;border-radius:10px;color:#f0f1f5;font-size:14px;margin-bottom:14px;outline:none;font-family:inherit;transition:border-color .2s}}
.pw-input:focus{{border-color:#818cf8;box-shadow:0 0 0 3px rgba(129,140,248,.15)}}
.pw-btn{{width:100%;padding:14px;background:#818cf8;color:#fff;border:none;border-radius:10px;font-size:14px;font-weight:700;cursor:pointer;transition:background .2s,transform .1s;font-family:inherit}}
.pw-btn:hover{{background:#6366f1}}.pw-btn:active{{transform:scale(.98)}}
.pw-err{{color:#f87171;font-size:12px;margin-top:10px;display:none}}
</style></head><body>
<div class="login-box">
<h1><span>Recharge.com</span> Scanner</h1>
<p>Enter password to view the dashboard</p>
<input type="password" class="pw-input" id="pwInput" placeholder="Password" onkeydown="if(event.key==='Enter')unlock()" autofocus>
<button class="pw-btn" onclick="unlock()">Unlock Dashboard</button>
<div class="pw-err" id="pwErr">Incorrect password. Please try again.</div>
</div>
<script>const D="{enc_b64}";{decrypt_js}</script></body></html>"""

    fname = f"recharge_dashboard_{DATE}.html"
    with open(fname,"w",encoding="utf-8") as f: f.write(html)
    with open("index.html","w",encoding="utf-8") as f: f.write(html)
    print(f"OK -> {fname} + index.html"); return fname

# =============================================================================
# SECTION 11 - WORD DOCUMENT
# =============================================================================

def _shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'),color)
    cell._tc.get_or_add_tcPr().append(s)

def build_docx(cands, ai, events, all_sig, top3=None):
    print("\nWORD DOCUMENT...", end=" ")
    doc = Document(); ex = ai.get("executive",{}); opps = ai.get("opportunities",[])
    comp_intel_data = ai.get("competitor_intel",{}); top3 = top3 or {"GMG":[],"ENT":[],"PPM":[],"MTU":[]}
    comp_intel = comp_intel_data.get("competitors",[]) if isinstance(comp_intel_data, dict) else comp_intel_data
    n_sources = len([k for k,v in all_sig.items() if len(v)>0])
    t = doc.add_heading("Recharge.com Opportunity Report",0); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{DATE} {TIME} | {n_sources} sources | 4-pass AI (grounded)"); r.font.color.rgb = RGBColor(128,128,128); r.font.size = Pt(10)
    # Confidential notice
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENTIAL - FOR INTERNAL USE ONLY"); r.font.size = Pt(8); r.font.color.rgb = RGBColor(180,180,180); r.font.all_caps = True
    # Horizontal rule
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '4B5563')
    pBdr.append(bottom); pPr.append(pBdr)
    # Table of Contents
    p = doc.add_paragraph(); r = p.add_run("TABLE OF CONTENTS"); r.bold = True; r.font.size = Pt(12)
    for section_name in ["Executive Summary", "Top 3 Per Category", "Competitor Intelligence", "Top Opportunities", "Events Calendar", "Predictions & Risks"]:
        doc.add_paragraph(f"  {section_name}", style="List Number")
    doc.add_page_break()
    doc.add_heading("Executive Summary",level=1); doc.add_paragraph(ex.get("summary",""))
    # Top 3 Per Category
    doc.add_heading("Top 3 Per Business Category",level=1)
    for bc_code in ["GMG","ENT","PPM","MTU"]:
        bc_name = BIZ_CAT_NAMES[bc_code]
        items = top3.get(bc_code,[])
        doc.add_heading(f"{bc_name} ({bc_code})",level=2)
        if not items:
            doc.add_paragraph("No items this week."); continue
        tbl = doc.add_table(rows=1,cols=4); tbl.style = "Table Grid"
        for i,h in enumerate(["#","Title","Score","Why Now"]):
            c2 = tbl.rows[0].cells[i]; c2.text = h; c2.paragraphs[0].runs[0].bold = True
        for i,item in enumerate(items,1):
            row = tbl.add_row().cells
            row[0].text = str(i); row[1].text = item.get("title","")[:50]
            row[2].text = str(item.get("score",0)); row[3].text = item.get("why_now","")[:50]
    # Competitor Intelligence
    doc.add_page_break()
    doc.add_heading("Competitor Intelligence",level=1)
    if isinstance(comp_intel_data, dict):
        ms = comp_intel_data.get("market_summary","")
        bt = comp_intel_data.get("biggest_threat","")
        og = comp_intel_data.get("opportunity_gap","")
        if ms: p = doc.add_paragraph(); p.add_run("Market Landscape: ").bold = True; p.add_run(ms)
        if bt: p = doc.add_paragraph(); p.add_run("Biggest Threat: ").bold = True; p.add_run(bt)
        if og: p = doc.add_paragraph(); p.add_run("Opportunity Gap: ").bold = True; p.add_run(og)
    if comp_intel:
        for ci in comp_intel:
            tl = ci.get('threat_level','').upper()
            ct = ci.get('type','').upper()
            doc.add_heading(f"{ci.get('name','')} [{ct}] - {tl} THREAT",level=2)
            doc.add_paragraph(ci.get("headline",""))
            if ci.get("strategy_signal"):
                p = doc.add_paragraph(); p.add_run("Strategy Signal: ").bold = True; p.add_run(ci.get("strategy_signal",""))
            for act in ci.get("activities",[])[:5]:
                doc.add_paragraph(act, style="List Bullet")
            if ci.get("news_highlights"):
                p = doc.add_paragraph(); p.add_run("In the News: ").bold = True
                for n in ci.get("news_highlights",[])[:3]: doc.add_paragraph(n, style="List Bullet")
            if ci.get("website_moves"):
                p = doc.add_paragraph(); p.add_run("Website: ").bold = True; p.add_run(ci.get("website_moves",""))
            if ci.get("notable"):
                p = doc.add_paragraph(); p.add_run("Notable: ").bold = True; p.add_run(ci.get("notable",""))
            if ci.get("our_response"):
                p = doc.add_paragraph(); p.add_run("Our Move: ").bold = True; p.add_run(ci.get("our_response",""))
    doc.add_heading("Top Opportunities",level=1)
    tbl = doc.add_table(rows=1,cols=7); tbl.style = "Table Grid"
    for i,h in enumerate(["#","Opportunity","Category","Score","Src","Urgency","Revenue Signal"]):
        c2 = tbl.rows[0].cells[i]; c2.text = h; c2.paragraphs[0].runs[0].bold = True
        _shade(c2,"1a237e"); c2.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    for i,o in enumerate(opps[:15],1):
        row = tbl.add_row().cells; mc = _match_cand(o.get("title",""),cands)
        row[0].text = str(i); row[1].text = o.get("title","")[:45]; row[2].text = o.get("category","")
        row[3].text = str(mc.score) if mc else "-"; row[4].text = str(mc.sources) if mc else "-"
        row[5].text = o.get("urgency",""); row[6].text = o.get("revenue_signal","")[:60]
        u = o.get("urgency","")
        if u == "critical": _shade(row[5],"ffcdd2")
        elif u == "high": _shade(row[5],"fff9c4")
    doc.add_heading("Events (next 60 days)",level=1)
    tbl = doc.add_table(rows=1,cols=4); tbl.style = "Table Grid"
    for i,h in enumerate(["Event","Category","Status","Details"]):
        c2 = tbl.rows[0].cells[i]; c2.text = h; c2.paragraphs[0].runs[0].bold = True
        _shade(c2,"1a237e"); c2.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    for e in events[:20]:
        row = tbl.add_row().cells; row[0].text = e["name"][:50]; row[1].text = e["category"]
        row[2].text = e["status"]; row[3].text = e["description"][:45]
        if e.get("urgency")=="critical": _shade(row[2],"ffcdd2")
        elif e.get("urgency")=="high": _shade(row[2],"fff9c4")
    for title,key in [("Predictions","predictions"),("Risk Watchlist","risks")]:
        items = ex.get(key,[])
        if items:
            doc.add_heading(title,level=1)
            for item in items: doc.add_paragraph(item,style="List Bullet")
    fname = f"recharge_report_v4_{DATE}.docx"; doc.save(fname); print(f"OK -> {fname}"); return fname

# =============================================================================
# SECTION 12 - EMAIL
# =============================================================================

def build_email_html(cands, ai, events, all_sig, top3=None, dashboard_url="", newsletter_copy=None):
    ex = ai.get("executive",{}); opps = ai.get("opportunities",[])
    comp_intel_data = ai.get("competitor_intel",{})
    comp_intel = comp_intel_data.get("competitors",[]) if isinstance(comp_intel_data, dict) else comp_intel_data
    n_sources = len([k for k,v in all_sig.items() if len(v)>0])
    total_sig = sum(len(v) for v in all_sig.values())
    top3 = top3 or {"GMG":[],"ENT":[],"PPM":[],"MTU":[]}
    nc = newsletter_copy or {}

    # The Big Picture
    big_picture = esc(nc.get("big_picture", ex.get("summary","")))

    # What's Hot section
    hot_items = nc.get("hot_items", [])
    if not hot_items:
        hot_items = [{"title":o.get("title","")[:40],"blurb":o.get("why_now","")[:60]} for o in opps[:3]]
    hot_html = ""
    for i, item in enumerate(hot_items[:3], 1):
        hot_html += f'''<table width="100%" cellpadding="0" cellspacing="0" style="margin-bottom:10px"><tr>
<td style="background:linear-gradient(135deg,#818cf811,#818cf805);border-left:3px solid #818cf8;padding:12px 16px;border-radius:0 8px 8px 0">
<p style="margin:0;font-weight:700;font-size:14px;color:#1f2937">{i}. {esc(item.get("title",""))}</p>
<p style="margin:4px 0 0;font-size:12px;color:#6b7280">{esc(item.get("blurb",""))}</p>
</td></tr></table>'''

    # Category Spotlights
    cat_blurbs = nc.get("category_blurbs", {})
    cat_html = ""
    for bc_code in ["GMG","ENT","PPM","MTU"]:
        bc_name = BIZ_CAT_NAMES[bc_code]; bc_emoji = BIZ_CAT_EMOJI[bc_code]
        bc_color = BIZ_CAT_COLORS[bc_code]
        items = top3.get(bc_code, [])
        blurbs = cat_blurbs.get(bc_code, [])
        items_html = ""
        for i, item in enumerate(items[:3], 1):
            blurb = ""
            if i <= len(blurbs): blurb = blurbs[i-1].get("blurb","")
            if not blurb: blurb = item.get("why_now","")[:60]
            title = item.get("title","")[:50]
            if i <= len(blurbs) and blurbs[i-1].get("title"):
                title = blurbs[i-1]["title"][:50]
            url = item.get("url","")
            title_html = f'<a href="{esc(url)}" style="color:#1f2937;text-decoration:none;font-weight:600">{esc(title)}</a>' if url else f'<span style="font-weight:600">{esc(title)}</span>'
            items_html += f'<p style="margin:0 0 10px;font-size:13px"><span style="color:{bc_color};font-weight:700">{i}.</span> {title_html}<br><span style="color:#6b7280;font-size:12px">{esc(blurb)}</span></p>'
        if not items_html:
            items_html = '<p style="margin:0;font-size:12px;color:#9ca3af">No updates this week.</p>'
        cat_html += f'''<table width="100%" cellpadding="0" cellspacing="0" style="border:1px solid #e5e7eb;border-radius:12px;overflow:hidden;margin-bottom:14px">
<tr><td style="background:{bc_color};padding:10px 16px"><p style="margin:0;color:#ffffff;font-size:13px;font-weight:700">{bc_emoji} {esc(bc_name)} ({bc_code})</p></td></tr>
<tr><td style="padding:14px 16px">{items_html}</td></tr></table>'''

    # Competitor Watch
    comp_watch = nc.get("competitor_watch", "")
    if not comp_watch and comp_intel:
        comp_watch = "; ".join(f"{ci.get('name','')}: {ci.get('headline','')}" for ci in comp_intel[:3])
    if not comp_watch:
        comp_watch = "No major competitor moves detected this week."

    # Don't Miss This (upcoming events)
    dont_miss = nc.get("dont_miss", [])
    if not dont_miss:
        dont_miss = [{"event":e["name"],"why_care":e.get("description","")} for e in events[:5] if e.get("urgency") in ("critical","high")]
    dm_html = ""
    for dm in dont_miss[:4]:
        dm_html += f'<tr><td style="padding:8px 12px;border-bottom:1px solid #f3f4f6;font-size:13px;font-weight:600;color:#1f2937">{esc(dm.get("event",""))}</td><td style="padding:8px 12px;border-bottom:1px solid #f3f4f6;font-size:12px;color:#6b7280">{esc(dm.get("why_care",""))}</td></tr>'

    # Crystal Ball
    crystal = nc.get("crystal_ball", "")
    if not crystal:
        preds = ex.get("predictions",[])
        crystal = " ".join(preds[:2]) if preds else "Stay tuned for next week."

    # Dashboard button
    dash_btn = f'''<tr><td style="padding:28px 32px;text-align:center">
<a href="{esc(dashboard_url)}" style="display:inline-block;background:linear-gradient(135deg,#4f46e5,#7c3aed);color:#ffffff;padding:16px 44px;border-radius:10px;text-decoration:none;font-weight:700;font-size:15px;box-shadow:0 4px 12px rgba(79,70,229,.3)">View Full Dashboard &rarr;</a>
</td></tr>''' if dashboard_url else ""

    return f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<style>
@media only screen and (max-width:480px) {{
  .email-container {{ width:100% !important; }}
  .email-padding {{ padding-left:16px !important; padding-right:16px !important; }}
  .cat-card {{ display:block !important; width:100% !important; }}
  .hide-mobile {{ display:none !important; }}
  h1 {{ font-size:18px !important; }}
  h2 {{ font-size:14px !important; }}
}}
</style></head>
<body style="margin:0;padding:0;background:#f8fafc;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Arial,sans-serif">
<!-- Preheader text -->
<div style="display:none;max-height:0;overflow:hidden;mso-hide:all">{esc(nc.get("big_picture", ex.get("summary",""))[:80])} &zwnj;&nbsp;&zwnj;&nbsp;</div>
<table width="100%" cellpadding="0" cellspacing="0" style="background:#f8fafc;padding:20px 0"><tr><td align="center">
<table class="email-container" width="100%" cellpadding="0" cellspacing="0" style="max-width:640px;width:100%;background:#ffffff;border-radius:12px;overflow:hidden;box-shadow:0 2px 8px rgba(0,0,0,0.06)">

<!-- HEADER -->
<tr><td style="background:linear-gradient(135deg,#0f172a,#1e1b4b,#312e81);padding:36px 32px" class="email-padding">
<table width="100%"><tr>
<td><h1 style="margin:0;color:#ffffff;font-size:22px;font-weight:800;letter-spacing:-0.02em">\U0001F680 This Week @ Recharge</h1>
<p style="margin:8px 0 0;color:#94a3b8;font-size:12px;font-weight:500">{DATE} &bull; {n_sources} sources &bull; {total_sig} signals</p></td>
<td style="text-align:right;vertical-align:top"><div style="background:rgba(129,140,248,.15);border-radius:8px;padding:8px 14px;display:inline-block">
<span style="font-size:24px;font-weight:800;color:#818cf8">{len([o for o in opps if o.get('urgency') in ('critical','high')])}</span><br>
<span style="font-size:10px;color:#94a3b8;text-transform:uppercase">Action Items</span></div></td>
</tr></table></td></tr>

<!-- QUICK STATS BAR -->
<tr><td style="background:#f0f0ff;padding:14px 32px" class="email-padding">
<table width="100%" cellpadding="0" cellspacing="0"><tr>
<td style="text-align:center;width:25%"><span style="font-size:20px;font-weight:800;color:#4f46e5">{total_sig}</span><br><span style="font-size:10px;color:#6b7280;text-transform:uppercase">Signals</span></td>
<td style="text-align:center;width:25%"><span style="font-size:20px;font-weight:800;color:#059669">{len(opps)}</span><br><span style="font-size:10px;color:#6b7280;text-transform:uppercase">Opportunities</span></td>
<td style="text-align:center;width:25%"><span style="font-size:20px;font-weight:800;color:#d97706">{len(events)}</span><br><span style="font-size:10px;color:#6b7280;text-transform:uppercase">Events</span></td>
<td style="text-align:center;width:25%"><span style="font-size:20px;font-weight:800;color:#dc2626">{len(ex.get('actions',[]))}</span><br><span style="font-size:10px;color:#6b7280;text-transform:uppercase">Actions</span></td>
</tr></table></td></tr>

<!-- THE BIG PICTURE -->
<tr><td style="padding:24px 28px" class="email-padding">
<table width="100%"><tr><td style="border-left:4px solid #4f46e5;padding-left:14px">
<h2 style="margin:0;font-size:15px;font-weight:700;color:#0f172a">\U0001F30D The Big Picture</h2>
<p style="margin:2px 0 0;font-size:11px;color:#64748b">Your weekly intelligence brief</p>
</td></tr></table>
<p style="margin:12px 0 0;font-size:14px;line-height:1.7;color:#374151">{big_picture}</p>
</td></tr>

<!-- Divider -->
<tr><td style="padding:0 28px"><div style="border-bottom:1px solid #e5e7eb;margin:4px 0"></div></td></tr>

<!-- WHAT'S HOT -->
<tr><td style="padding:20px 28px 16px" class="email-padding">
<table width="100%"><tr><td style="border-left:4px solid #d97706;padding-left:14px">
<h2 style="margin:0;font-size:15px;font-weight:700;color:#0f172a">\U0001F525 What's Hot This Week</h2>
<p style="margin:2px 0 0;font-size:11px;color:#64748b">The headlines that matter for Recharge.com</p>
</td></tr></table>
{hot_html}
</td></tr>

<!-- CATEGORY SPOTLIGHTS -->
<tr><td style="padding:0 32px 20px">
<h2 style="margin:0 0 14px;font-size:16px;color:#1f2937">\U0001F4CA Category Spotlights</h2>
{cat_html}
</td></tr>

<!-- COMPETITOR WATCH -->
<tr><td style="padding:0 32px 20px">
<h2 style="margin:0 0 8px;font-size:16px;color:#1f2937">\U0001F50D Competitor Watch</h2>
<table width="100%" cellpadding="0" cellspacing="0" style="background:#fef3c7;border-radius:8px;overflow:hidden">
<tr><td style="padding:14px 16px;font-size:13px;color:#92400e;line-height:1.6">{esc(comp_watch)}</td></tr>
</table>
</td></tr>

<!-- DON'T MISS THIS -->
<tr><td style="padding:0 32px 20px">
<h2 style="margin:0 0 8px;font-size:16px;color:#1f2937">\U0001F4C5 Don't Miss This</h2>
<table width="100%" cellpadding="0" cellspacing="0" style="border:1px solid #e5e7eb;border-radius:8px;overflow:hidden">
<tr style="background:#f9fafb"><td style="padding:8px 12px;font-size:10px;color:#6b7280;text-transform:uppercase;font-weight:600;border-bottom:1px solid #e5e7eb">Event</td>
<td style="padding:8px 12px;font-size:10px;color:#6b7280;text-transform:uppercase;font-weight:600;border-bottom:1px solid #e5e7eb">Why You Should Care</td></tr>
{dm_html}</table>
</td></tr>

<!-- CRYSTAL BALL -->
<tr><td style="padding:0 32px 20px">
<h2 style="margin:0 0 8px;font-size:16px;color:#1f2937">\U0001F52E Crystal Ball</h2>
<table width="100%" cellpadding="0" cellspacing="0" style="background:#ede9fe;border-radius:8px;overflow:hidden">
<tr><td style="padding:14px 16px;font-size:13px;color:#5b21b6;line-height:1.6">{esc(crystal)}</td></tr>
</table>
</td></tr>

<!-- DASHBOARD CTA -->
{dash_btn}

<!-- FOOTER -->
<tr><td style="padding:24px 32px;background:#0f172a;text-align:center">
<p style="margin:0;font-size:13px;font-weight:700;color:#818cf8">Recharge.com Opportunity Scanner</p>
<p style="margin:6px 0 0;font-size:11px;color:#64748b">v5.0 &bull; {n_sources} sources &bull; {total_sig} signals &bull; AI-powered weekly intelligence</p>
<p style="margin:8px 0 0;font-size:10px;color:#475569">Powered by Google Gemini + real-time data. Open the dashboard for the full picture.</p>
</td></tr>

</table></td></tr></table></body></html>"""

def send_email(html_body, subject=None):
    to_addr = os.environ.get("EMAIL_TO","").strip()
    from_addr = os.environ.get("EMAIL_FROM","").strip()
    smtp_host = os.environ.get("SMTP_HOST","smtp.gmail.com").strip()
    smtp_port_raw = os.environ.get("SMTP_PORT","587").strip()
    smtp_user = os.environ.get("SMTP_USER","").strip()
    smtp_pass = os.environ.get("SMTP_PASS","").strip().replace(" ","")  # Gmail app passwords: strip spaces

    # Debug: show what we have (mask password)
    print(f"EMAIL CONFIG: to={to_addr!r}, from={from_addr!r}, host={smtp_host!r}, port={smtp_port_raw!r}, user={smtp_user!r}, pass_len={len(smtp_pass)}")
    sys.stdout.flush()

    missing = [n for n,v in [("EMAIL_TO",to_addr),("EMAIL_FROM",from_addr),("SMTP_USER",smtp_user),("SMTP_PASS",smtp_pass)] if not v]
    if missing:
        print(f"EMAIL SEND... skipped (missing secrets: {', '.join(missing)})"); sys.stdout.flush()
        return False

    smtp_port = int(smtp_port_raw)
    print(f"EMAIL SEND... connecting to {smtp_host}:{smtp_port} as {smtp_user}"); sys.stdout.flush()
    import smtplib; from email.mime.multipart import MIMEMultipart; from email.mime.text import MIMEText
    if not subject: subject = f"This Week @ Recharge | {DATE} | Top Opportunities + Category Highlights"
    msg = MIMEMultipart('alternative'); msg['Subject'] = subject; msg['From'] = from_addr; msg['To'] = to_addr
    msg.attach(MIMEText(html_body,'html'))
    try:
        print(f"EMAIL SEND... opening SMTP connection..."); sys.stdout.flush()
        server = smtplib.SMTP(smtp_host, smtp_port, timeout=30)
        server.set_debuglevel(0)  # Set to 1 for debugging only
        server.ehlo()
        print(f"EMAIL SEND... STARTTLS..."); sys.stdout.flush()
        server.starttls()
        server.ehlo()
        print(f"EMAIL SEND... logging in..."); sys.stdout.flush()
        server.login(smtp_user, smtp_pass)
        print(f"EMAIL SEND... sending to {to_addr}..."); sys.stdout.flush()
        recipients = [a.strip() for a in to_addr.split(",")]
        server.sendmail(from_addr, recipients, msg.as_string())
        server.quit()
        print(f"EMAIL SEND... OK -> {to_addr}"); sys.stdout.flush(); return True
    except Exception as e:
        print(f"EMAIL SEND... FAILED: {type(e).__name__}: {e}"); sys.stdout.flush()
        import traceback; traceback.print_exc(); sys.stdout.flush()
        return False

# =============================================================================
# SECTION 13 - EVENTS CALENDAR
# =============================================================================

def get_events():
    today = NOW
    cal = [
        (2,27,28,"Resident Evil Requiem","Game Release","PS5/XSX/Switch 2/PC",9),
        (3,19,20,"Crimson Desert","Game Release","PS5/XSX/PC",8),
        (3,20,21,"Saros (Housemarque)","PlayStation","PS5 exclusive",8),
        (3,27,28,"007 First Light","Game Release","James Bond",8),
        (3,15,31,"Marathon Launch","PlayStation","Bungie extraction shooter",9),
        (8,15,16,"Madden NFL 27","Game Release","Annual",8),
        (9,1,7,"NBA 2K27","Game Release","Annual",8),
        (9,25,30,"EA Sports FC 27","EA Sports FC","New FC game",10),
        (10,15,31,"Call of Duty 2026","Call of Duty","First on Switch 2",10),
        (11,19,19,"GTA 6 LAUNCH","GTA","Biggest release in history",10),
        (6,1,30,"Fable","Xbox","Xbox exclusive RPG",9),
        (10,1,31,"Marvel's Wolverine","PlayStation","Insomniac PS5",9),
        (12,1,31,"Forza Horizon 6","Xbox","Set in Japan",8),
        (1,22,22,"Xbox Developer Direct","Xbox","Confirmed",9),
        (6,7,8,"Xbox Games Showcase","Xbox","Post-SGF",9),
        (1,1,7,"Game Pass Wave 1","Xbox","New additions",7),
        (2,1,7,"Game Pass Feb","Xbox","New additions",7),
        (3,1,7,"Game Pass Mar","Xbox","New additions",7),
        (2,15,28,"State of Play","PlayStation","Feb broadcast",8),
        (5,25,31,"PlayStation Showcase","PlayStation","Major",9),
        (9,15,25,"State of Play Fall","PlayStation","Sep",8),
        (1,1,7,"PS Plus Jan","PlayStation","Monthly free games",8),
        (2,1,7,"PS Plus Feb","PlayStation","Monthly free games",8),
        (3,1,7,"PS Plus Mar","PlayStation","Monthly free games",8),
        (4,1,7,"PS Plus Apr","PlayStation","Monthly free games",8),
        (5,1,7,"PS Plus May","PlayStation","Monthly free games",8),
        (6,1,7,"PS Plus Jun","PlayStation","Monthly free games",8),
        (2,10,20,"Nintendo Direct Feb","Nintendo","Annual pattern",8),
        (6,10,15,"Nintendo Direct Summer","Nintendo","SGF period",9),
        (9,10,20,"Nintendo Direct Fall","Nintendo","Sep pattern",8),
        (2,23,23,"Steam Next Fest","Steam","Feb 23-Mar 2",7),
        (3,19,19,"Steam Spring Sale","Steam","Mar 19-26",9),
        (6,26,26,"Steam Summer Sale","Steam","Biggest sale",10),
        (11,25,25,"Steam Autumn Sale","Steam","Pre-BF",9),
        (12,18,18,"Steam Winter Sale","Steam","Through Jan 2",10),
        (1,12,19,"Steam Detective Fest","Steam","Mystery games",8),
        (2,9,16,"Steam PvP Fest","Steam","Competitive",8),
        (4,20,27,"Steam Medieval Fest","Steam","Knights & castles",8),
        (5,4,11,"Steam Deckbuilders Fest","Steam","Card strategy",8),
        (8,3,10,"Steam Cyberpunk Fest","Steam","Neon dystopian",8),
        (10,19,26,"Steam Next Fest Oct","Steam","Fall demos",9),
        (10,26,31,"Steam Scream V","Steam","Halloween horror",9),
        (1,10,25,"EA FC TOTY","EA Sports FC","Team of the Year",10),
        (2,5,20,"EA FC Future Stars","EA Sports FC","Young talents",8),
        (3,15,31,"FUT Birthday","EA Sports FC","Anniversary",8),
        (5,1,31,"EA FC TOTS","EA Sports FC","Team of the Season",9),
        (6,15,30,"EA FC Futties","EA Sports FC","End of cycle",8),
        (11,20,30,"EA FC Black Friday","EA Sports FC","Lightning rounds",9),
        (3,9,13,"GDC","Esports","San Francisco",7),
        (3,26,29,"PAX East","Esports","Boston",7),
        (4,27,30,"iicon (E3 successor)","Esports","Las Vegas",9),
        (6,5,8,"Summer Game Fest","Esports","LA",10),
        (8,26,30,"Gamescom","Esports","Cologne",9),
        (9,12,13,"BlizzCon","Esports","Anaheim",8),
        (9,17,21,"Tokyo Game Show","Esports","Chiba",8),
        (12,5,12,"The Game Awards","Esports","Major reveals",9),
        (7,6,6,"Esports World Cup","Esports","Riyadh $70M+",9),
        (10,1,31,"LoL Worlds","Esports","NYC",9),
        (1,14,14,"Genshin 6.3","Genshin Impact","Lantern Rite",9),
        (2,25,25,"Genshin 6.4","Genshin Impact","Varka banner",9),
        (4,8,8,"Genshin 6.5","Genshin Impact","Hexenzirkel",8),
        (8,12,12,"Genshin 7.0","Genshin Impact","Anniversary",10),
        (2,14,14,"HSR 4.0","Honkai","Planarcadia",9),
        (4,26,30,"HSR 2nd Anniversary","Honkai","Major rewards",9),
        (12,1,4,"Spotify Wrapped","Spotify","Viral",10),
        (11,12,12,"Disney+ Day","Disney Plus","Annual",8),
        (5,22,22,"Mandalorian & Grogu","Disney Plus","Star Wars film",9),
        (11,26,26,"Stranger Things S5","Netflix","Final season",9),
        (12,18,18,"Avengers: Doomsday","Gift Cards","Biggest MCU",10),
        (2,17,17,"Chinese New Year","Gift Cards","Fire Horse",9),
        (7,7,10,"Amazon Prime Day","Gift Cards","Confirmed",9),
        (11,11,11,"Singles Day","Gift Cards","World's largest",9),
        (11,27,27,"Black Friday","Gift Cards","Peak sales",10),
        (11,30,30,"Cyber Monday","Gift Cards","Digital focus",9),
        (2,1,14,"Valentine's Day","Gift Cards","Gift peak",8),
        (5,1,12,"Mother's Day","Gift Cards","Major gift event",8),
        (6,1,15,"Father's Day","Gift Cards","Gaming gifting",8),
        (12,1,24,"Christmas Season","Gift Cards","Peak buying",10),
        (1,1,15,"Fortnite Winterfest","Fortnite","Holiday event",8),
        (10,15,31,"Fortnitemares","Fortnite","Halloween",8),
        (1,1,7,"Winter Anime","Crunchyroll","New premieres",7),
        (4,1,7,"Spring Anime","Crunchyroll","New premieres",7),
        (7,1,7,"Summer Anime","Crunchyroll","New premieres",7),
        (10,1,7,"Fall Anime","Crunchyroll","New premieres",7),
        (2,1,1,"Grammy Awards","Spotify","Music night",9),
        (3,15,15,"Oscars","Netflix","Film night",9),
        (3,21,21,"Monster Hunter Wilds","Game Release","Capcom PC/PS5/XSX",10),
        (2,28,28,"Elden Ring Nightreign","Game Release","FromSoftware",9),
    ]
    events = []
    for mo,d1,d2,name,cat,desc,pri in cal:
        try:
            s = datetime(today.year,mo,d1); e = datetime(today.year,mo,d2)
            if s < today-timedelta(days=30): s = datetime(today.year+1,mo,d1); e = datetime(today.year+1,mo,d2)
            days = (s-today).days
            if -10<=days<=60:
                if days<=0 and (e-today).days>=0: st,urg = "ACTIVE NOW","critical"
                elif days<=7: st,urg = f"{days}d","high"
                elif days<=14: st,urg = f"{days}d","medium"
                else: st,urg = f"{days}d","low"
                events.append({"name":name,"category":cat,"description":desc,"status":st,"urgency":urg,"days_until":days,"priority":pri,"is_live":False})
        except Exception as e:
            log.debug(f"Event '{name}': {e}"); continue
    seen = set()
    for q in ["game announcement today","release date announced","new update live","free games announced",
              "PlayStation Plus reveal","Game Pass announced","Nintendo Direct date","Steam sale date",
              "EA FC promo","Genshin banner","anime premiere","Crunchyroll new"]:
        try:
            feed = feedparser.parse(f"https://news.google.com/rss/search?q={quote(q)}+when:3d&hl=en-US&gl=US&ceid=US:en")
            for e in feed.entries[:3]:
                t = e.get("title",""); src = "News"
                if " - " in t: t,src = t.rsplit(" - ",1)
                k = t[:40].lower()
                if k in seen: continue
                seen.add(k); lo = t.lower()
                if any(x in lo for x in ["announce","reveal","launch","release","live now","available now","sale","event","promo","premiere"]):
                    cc = cats(t); cat2 = cc[0] if cc else "Gaming"
                    if any(x in lo for x in ["live now","out now","available now","today"]): urg,st = "critical","LIVE NOW"
                    elif any(x in lo for x in ["tomorrow","coming soon","this week"]): urg,st = "high","SOON"
                    else: urg,st = "medium","ANNOUNCED"
                    events.append({"name":t[:80],"category":cat2,"description":f"via {src}","status":st,"urgency":urg,
                                   "days_until":0 if urg=="critical" else 1,"priority":9 if urg=="critical" else 7,"is_live":True})
            time.sleep(0.2)
        except Exception as e:
            log.debug(f"Event discovery '{q[:30]}': {e}"); continue
    for svc in ["Netflix","Prime Video","Disney Plus","Crunchyroll"]:
        try:
            feed = feedparser.parse(f"https://news.google.com/rss/search?q={quote(svc)}+new+release+when:7d&hl=en-US&gl=US&ceid=US:en")
            for e in feed.entries[:5]:
                t = e.get("title","")
                if " - " in t: t = t.rsplit(" - ",1)[0]
                if any(k in t.lower() for k in ["premieres","launches","releases","arrives","streaming","drops"]):
                    events.append({"name":f"{svc}: {t[:50]}","category":svc,"description":"Streaming release","status":"NOW","urgency":"high","days_until":0,"priority":7,"is_live":True})
            time.sleep(0.1)
        except Exception as e:
            log.debug(f"Streaming events {svc}: {e}"); continue
    return sorted(events,key=lambda x:(-x["priority"] if x["urgency"]=="critical" else 0,x.get("days_until",99)))

# =============================================================================
# MAIN
# =============================================================================

def main():
    t0 = time.time()
    _phase_times = {}
    def _phase(name): _phase_times[name] = time.time()
    def _phase_end(name):
        if name in _phase_times:
            elapsed = time.time() - _phase_times[name]
            log.info(f"  (took {elapsed:.1f} seconds)")
            return elapsed
        return 0

    log.info("")
    log.info("=" * 70)
    log.info("  Hi! I'm the Recharge.com Opportunity Scanner (v5.0).")
    log.info(f"  Today is {DATE}, and I'm starting my scan at {TIME}.")
    log.info("=" * 70)
    log.info("")
    log.info("My job is to scan 16+ data sources across the internet and find sales")
    log.info("opportunities for Recharge.com — gift cards, game credits, streaming")
    log.info("subscriptions, and mobile top-up. Let me get to work!")
    log.info("")
    print("="*60); print("  RECHARGE.COM OPPORTUNITY SCANNER v5.0"); print(f"  {DATE} {TIME}"); print("="*60)
    sys.stdout.flush()

    # --- Config summary ---
    log.info("First, let me check my setup...")
    log.info(f"  AI brain (Gemini): {'ready to go!' if GEMINI_KEY else 'not configured — I will skip AI analysis'}")
    log.info(f"  News search (Oxylabs): {'ready — I can search Google News' if OXYLABS_USER and OXYLABS_PASS else 'not configured — I will rely on RSS feeds only'}")
    log.info(f"  Dashboard password: {'enabled — only people with the password can view it' if DASH_PASSWORD else 'disabled — anyone with the link can see it'}")
    log.info(f"  Competitors I'm watching: {', '.join(list(COMPETITORS.keys()) + [c for c in SITEMAP_COMPETITORS if c not in COMPETITORS])}")
    log.info(f"  Competitor blogs I read: {', '.join(BLOG_OVERRIDES.keys())}")
    log.info(f"  Running on: {'Google Colab' if IS_COLAB else 'local machine / GitHub Actions'}")

    # --- Phase 1: Events ---
    log.info("")
    log.info("=" * 50)
    log.info("STEP 1: Let me check what events and game releases are coming up...")
    log.info("=" * 50)
    _phase("events")
    events = get_events()
    _phase_end("events")
    log.info(f"  I found {len(events)} upcoming events — game launches, sales, streaming releases, and more.")
    for ev in events[:10]:
        log.info(f"    - {ev.get('name','')[:70]} ({ev.get('category','')})")
    if len(events) > 10:
        log.info(f"    ... and {len(events)-10} more")
    print(f"Events: {len(events)}"); sys.stdout.flush()

    # --- Phase 2: Data Fetching ---
    log.info("")
    log.info("=" * 50)
    log.info("STEP 2: Now I'll scan all my data sources for opportunities...")
    log.info("=" * 50)
    log.info("  I'm about to check 16+ sources at the same time — Steam, Reddit,")
    log.info("  Google News (via Oxylabs), competitor websites, game deal sites, etc.")
    log.info("  Each source runs in its own thread so I can be fast about it.")
    log.info("")
    _phase("fetch")
    all_sig = fetch_all()
    _phase_end("fetch")

    # Detailed per-source log
    _source_labels = {
        "steam":"Steam Top Sellers","reddit":"Reddit Gaming Posts","news":"News RSS Feeds",
        "oxylabs_news":"Google News (Oxylabs)","competitor":"Competitor Homepages",
        "sitemap":"Competitor Blogs","trends":"Google Search Trends",
        "cheapshark":"CheapShark Game Deals","steamspy":"SteamSpy Player Data",
        "gamerpower":"GamerPower Free Games","epic":"Epic Games Free Games",
        "gog":"GOG Popular Games","humble":"Humble Bundle","freetogame":"Free-to-Play Games",
        "anime":"Top Anime (Jikan)","youtube":"YouTube Gaming Channels","wiki":"Wikipedia Page Activity",
    }
    log.info("")
    log.info("Here's what I found from each source:")
    total_signals = 0
    for src_name in sorted(all_sig.keys()):
        sigs = all_sig[src_name]
        total_signals += len(sigs)
        label = _source_labels.get(src_name, src_name)
        if sigs:
            log.info(f"  {label}: {len(sigs)} items")
            for s in sigs[:3]:
                log.info(f"      e.g. {s.title[:75]}")
        else:
            log.info(f"  {label}: came back empty (might be temporarily down)")
    log.info(f"")
    log.info(f"  In total, I collected {total_signals} raw data points.")

    # --- Phase 3: Score Normalization ---
    log.info("")
    log.info("=" * 50)
    log.info("STEP 3: Now I'll score and rank everything fairly...")
    log.info("=" * 50)
    log.info("  Each source uses different scoring systems, so I normalize them all")
    log.info("  to a 0-100 scale so they can be compared apples-to-apples.")
    _phase("normalize")
    normalize_scores(all_sig)
    _phase_end("normalize")

    # --- Phase 4: Dedup & Composite Scoring ---
    log.info("")
    log.info("=" * 50)
    log.info("STEP 4: Removing duplicates and calculating final scores...")
    log.info("=" * 50)
    log.info("  Many sources report the same news. I'll merge duplicates and give")
    log.info("  higher scores to items that appear across multiple sources.")
    _phase("dedup")
    cands = comp_score(dedup(all_sig))
    _phase_end("dedup")
    multi_source = len([c for c in cands if c.sources >= 2])
    log.info(f"")
    log.info(f"  After merging, I have {len(cands)} unique opportunities.")
    log.info(f"  Of these, {multi_source} appeared in 2+ sources — those are the strongest signals.")
    # Keyword category breakdown
    cat_counter = Counter()
    for c in cands:
        for cat in (c.categories if c.categories else ["Uncategorized"]):
            cat_counter[cat] += 1
    top_kw_cats = cat_counter.most_common(10)
    log.info(f"  Keyword categories found: {', '.join(f'{cat} ({n})' for cat, n in top_kw_cats)}")
    log.info(f"")
    log.info(f"  Here are the top 10 by score:")
    for i, c in enumerate(cands[:10], 1):
        multi_tag = f" [from {c.sources} sources]" if c.sources >= 2 else ""
        log.info(f"    {i:2d}. {c.title[:65]}{multi_tag} (score: {c.score:.0f})")

    # --- Step 5: History Comparison ---
    log.info("")
    log.info("=" * 50)
    log.info("STEP 5: Let me compare with last week's results...")
    log.info("=" * 50)
    log.info("  I'll check if anything is trending up or down compared to my last scan.")
    _phase("trends")
    prev_history = load_previous_history()
    trends = compute_trends(cands, all_sig, prev_history)
    _phase_end("trends")
    if trends["kpi_deltas"]:
        d = trends["kpi_deltas"]
        log.info(f"  I found a previous scan! Here's what changed:")
        sig_d = d.get("total_signals", 0)
        cand_d = d.get("candidates", 0)
        log.info(f"    Data points: {sig_d:+d} vs last week")
        log.info(f"    Unique opportunities: {cand_d:+d} vs last week")
        if trends.get("new_entries"):
            log.info(f"    Brand new items this week: {len(trends['new_entries'])}")
        if trends.get("dropped"):
            log.info(f"    Items that dropped off since last week: {len(trends['dropped'])}")
        print(f"  Week-over-week: signals {sig_d:+d}, candidates {cand_d:+d}")
    else:
        log.info(f"  No previous scan found — this is either my first run or the history file was removed.")

    # --- Step 6: AI Analysis ---
    log.info("")
    log.info("=" * 50)
    log.info("STEP 6: I'm now asking Google Gemini AI to review all the data...")
    log.info("=" * 50)
    log.info("  I'll have the AI make 4 passes over everything:")
    log.info("    I'll verify facts with Google Search grounding...")
    log.info("    I'll pick the most important priorities...")
    log.info("    I'll write the executive brief for leadership...")
    log.info("    I'll analyze competitors and their moves...")
    log.info("  This is the longest step — it may take a few minutes.")
    _phase("ai")
    ai = run_ai(cands, events, all_sig.get("competitor",[]), all_sig)
    _phase_end("ai")
    sys.stdout.flush()
    n_opps = len(ai.get('opportunities',[]))
    n_comp = len(ai.get('competitor_intel',[]))
    log.info(f"")
    log.info(f"  I found {n_opps} actionable opportunities.")
    log.info(f"  I analyzed {n_comp} competitor moves.")
    if ai.get('executive',{}).get('summary',''):
        log.info(f"  Executive summary: written ({len(ai['executive']['summary'])} characters).")
    for opp in ai.get("opportunities",[])[:5]:
        urg = opp.get('urgency','').upper()
        log.info(f"    - [{urg}] {opp.get('title','')[:65]}")

    # Compute top 3 per business category
    _biz_labels = {"GMG":"Gaming","ENT":"Entertainment","PPM":"Prepaid/Payments","MTU":"Mobile Top-Up"}
    top3 = get_top3_per_biz_cat(cands, ai.get("opportunities", []))
    log.info(f"")
    log.info(f"  Here are the best opportunities by business area:")
    for k,v in top3.items():
        label = _biz_labels.get(k, k)
        log.info(f"    {label} ({k}): {len(v)} top picks")
    print(f"\nTop 3 per category: " + ", ".join(f"{k}={len(v)}" for k,v in top3.items())); sys.stdout.flush()

    # --- Step 7: Newsletter ---
    log.info("")
    log.info("=" * 50)
    log.info("STEP 7: Now I'll write the email newsletter...")
    log.info("=" * 50)
    log.info("  I'm crafting a polished newsletter summary to email to the team.")
    _phase("newsletter")
    newsletter_copy = None
    try:
        newsletter_copy = pass_newsletter(
            ai.get("opportunities",[]), top3,
            ai.get("competitor_intel",[]), events, ai.get("executive",{}))
        log.info(f"  Newsletter text ready ({len(newsletter_copy or '')} characters)")
    except Exception as e:
        log.warning(f"  Newsletter writing failed: {e}")
        log.warning(f"  (The rest of the report will still be generated)")
        print(f"[WARN] Newsletter copy failed: {e}"); sys.stdout.flush()
    _phase_end("newsletter")

    # --- Google Sheets (optional) ---
    _phase("sheets")
    sheets_url = write_sheets(cands, ai, events); sys.stdout.flush()
    _phase_end("sheets")

    # --- Step 8: Build Outputs ---
    log.info("")
    log.info("=" * 50)
    log.info("STEP 8: Almost done! Let me create all the output files...")
    log.info("=" * 50)

    # HTML Dashboard
    _phase("html")
    log.info("  I'm building the interactive HTML dashboard...")
    print("\n[STEP] Building HTML..."); sys.stdout.flush()
    try:
        html_file = build_html(cands, ai, events, all_sig, top3, trends)
        size_kb = os.path.getsize(html_file)/1024
        log.info(f"  Dashboard created: {html_file} ({size_kb:.0f} KB)")
        print(f"[STEP] HTML done: {html_file}"); sys.stdout.flush()
    except Exception as e:
        log.error(f"  Dashboard creation failed: {e}")
        print(f"[ERROR] build_html failed: {type(e).__name__}: {e}"); sys.stdout.flush()
        import traceback; traceback.print_exc(); sys.stdout.flush()
        html_file = "index.html"
        with open(html_file,"w") as f: f.write("<h1>Scanner ran but HTML build failed</h1>")
    _phase_end("html")

    # Word Document
    _phase("docx")
    log.info("  I'm creating the Word report for offline sharing...")
    print("[STEP] Building Word doc..."); sys.stdout.flush()
    try:
        docx_file = build_docx(cands, ai, events, all_sig, top3)
        log.info(f"  Word report created: {docx_file}")
        print(f"[STEP] Word done: {docx_file}"); sys.stdout.flush()
    except Exception as e:
        log.error(f"  Word report creation failed: {e}")
        print(f"[ERROR] build_docx failed: {type(e).__name__}: {e}"); sys.stdout.flush()
        docx_file = "none"
    _phase_end("docx")

    # Email
    _phase("email")
    log.info("  I'm formatting the email newsletter...")
    print("[STEP] Building email..."); sys.stdout.flush()
    try:
        email_html = build_email_html(cands, ai, events, all_sig, top3, DASHBOARD_URL, newsletter_copy)
        email_file = f"recharge_email_{DATE}.html"
        with open(email_file,"w",encoding="utf-8") as f: f.write(email_html)
        log.info(f"  Email newsletter created: {email_file} ({len(email_html)/1024:.0f} KB)")
        print(f"[STEP] Email HTML done: {email_file}"); sys.stdout.flush()
    except Exception as e:
        log.error(f"  Email creation failed: {e}")
        print(f"[ERROR] build_email_html failed: {type(e).__name__}: {e}"); sys.stdout.flush()
        email_html = ""; email_file = "none"
    _phase_end("email")

    # Send email
    log.info("  I'm sending the newsletter to the team...")
    print("[STEP] Sending email..."); sys.stdout.flush()
    try:
        send_email(email_html)
        log.info("  Email sent successfully!")
    except Exception as e:
        log.warning(f"  Could not send email: {e}")
        log.warning(f"  (Email credentials may not be configured — the HTML file is still saved)")
        print(f"[WARN] send_email: {type(e).__name__}: {e}"); sys.stdout.flush()

    # Save weekly history for future trend comparison
    log.info("  I'm saving this week's results so I can compare next time...")
    try:
        history_file = f"history_{DATE}.json"
        history_data = {
            "date": DATE, "version": "5.0",
            "kpis": {"total_signals": sum(len(v) for v in all_sig.values()),
                     "candidates": len(cands),
                     "multi_source": len([c for c in cands if c.sources >= 2]),
                     "sources_active": len([k for k,v in all_sig.items() if len(v)>0])},
            "top_candidates": [{"title":c.title, "score":c.score, "sources":c.sources,
                                "category":c.category, "biz_category":c.biz_category}
                               for c in cands[:50]],
        }
        with open(history_file, "w") as f: json.dump(history_data, f, indent=2)
        log.info(f"  History saved: {history_file}")
    except Exception as e: log.debug(f"History save: {e}")

    # --- Final Summary ---
    elapsed = time.time()-t0; ex = ai.get("executive",{}); n_sources = len([k for k,v in all_sig.items() if len(v)>0])

    log.info("")
    log.info("=" * 70)
    log.info("  That's a wrap! Here's a summary of everything I did:")
    log.info("=" * 70)
    log.info(f"")
    log.info(f"  Total time:  {elapsed/60:.1f} minutes ({elapsed:.0f} seconds)")
    log.info(f"")
    log.info(f"  I checked {n_sources} out of {len(all_sig)} data sources.")
    log.info(f"  I collected {sum(len(v) for v in all_sig.values())} raw data points.")
    log.info(f"  After deduplication, I identified {len(cands)} unique opportunities.")
    log.info(f"  {multi_source} of those were confirmed by 2+ sources (strongest signals).")
    log.info(f"  I picked {len(ai.get('opportunities',[]))} top priorities for the team.")
    log.info(f"  I tracked {len(events)} upcoming events.")
    log.info(f"")
    log.info(f"  Files I created:")
    log.info(f"    Dashboard (HTML): {html_file}")
    log.info(f"    Email newsletter: {email_file}")
    log.info(f"    Word report:      {docx_file}")
    if sheets_url: log.info(f"    Google Sheet:     {sheets_url}")
    log.info(f"    This log file:    {_log_file}")
    log.info(f"")
    if ex.get('summary'):
        log.info(f"  My executive summary:")
        for line in ex['summary'].split('. '):
            if line.strip():
                log.info(f"    {line.strip()}.")
    log.info("")
    log.info("  Until next time! - The Recharge Scanner")
    log.info("=" * 70)

    print("\n" + "="*60)
    print(f"  DONE in {elapsed:.0f}s | {n_sources} sources active")
    print(f"  HTML:  {html_file}")
    print(f"  Email: {email_file}")
    print(f"  Word:  {docx_file}")
    if sheets_url: print(f"  Sheet: {sheets_url}")
    print(f"  Log:   {_log_file}")
    print("="*60)
    print(f"\n{ex.get('summary','')}")
    print("\nActions:")
    for i,a in enumerate(ex.get("actions",[]),1): print(f"  {i}. {a}")
    sys.stdout.flush()
    return html_file, docx_file

if __name__ == "__main__":
    try:
        html_f, docx_f = main()
    except Exception as e:
        print(f"\n[FATAL] {type(e).__name__}: {e}"); sys.stdout.flush()
        import traceback; traceback.print_exc()
        sys.exit(1)
    if IS_COLAB:
        try:
            from google.colab import files; files.download(html_f); files.download(docx_f)
        except ImportError: pass
